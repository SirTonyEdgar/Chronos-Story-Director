"""
Chronos Story Director - Backend Engine (Logic Layer)
=====================================================
The "Brain" of the operation.
- Handles all AI Processing (Scene Drafting, Analysis, Chat).
- Manages Workflows (LangGraph).
- Delegates persistence/storage to 'database_manager.py'.
"""

import sqlite3
import os
import re
import json
import shutil
import datetime
import glob
from io import BytesIO
from typing import TypedDict, Optional, List, Dict, Any
from dotenv import load_dotenv

# --- INTERNAL IMPORTS ---
try:
    from . import database_manager as db
except ImportError:
    import database_manager as db

# --- THIRD-PARTY DEPENDENCIES ---
from google import genai as new_genai
from langgraph.graph import StateGraph, END
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage
from fpdf import FPDF
from ebooklib import epub

try:
    from langchain_openai import ChatOpenAI
except ImportError:
    ChatOpenAI = None

try:
    from langchain_anthropic import ChatAnthropic
except ImportError:
    ChatAnthropic = None

# --- CONFIGURATION ---
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ENV_PATH = os.path.join(BASE_DIR, ".env")
load_dotenv(ENV_PATH)

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")

# --- TYPE DEFINITIONS ---
class StoryState(TypedDict):
    """
    State schema for the scene generation workflow (LangGraph).
    """
    profile_name: str
    chapter_num: Optional[int]
    part_num: Optional[int]
    year: int
    date_str: str
    time_str: str 
    scene_title: str
    scene_brief: str
    scene_outline: str
    timeline: str
    current_draft: str
    revision_count: int
    is_grounded: bool
    recent_context: str 
    banned_words: str
    use_fog_of_war: bool
    context_files: List[str]
    critique_notes: str
    style_notes: str
    style_result: str
    retrieved_ids: List[int]
    _planner_tokens: dict
    _drafter_tokens: dict
    _validator_tokens: dict
    _style_tokens: dict
    voice_notes: str
    voice_result: str
    pov_context: str

# ==========================================
# GENERATION STATUS TRACKER
# ==========================================
_generation_status: dict = {}

def update_generation_status(profile: str, step: int, step_name: str, message: str):
    """Updates the in-memory generation status for a profile."""
    _generation_status[profile] = {
        "step": step,
        "total_steps": 5,
        "step_name": step_name,
        "message": message,
        "active": True
    }

def clear_generation_status(profile: str):
    """Clears the generation status for a profile."""
    _generation_status[profile] = {"active": False}

def get_generation_status(profile: str) -> dict:
    """Returns the current generation status for a profile."""
    return _generation_status.get(profile, {"active": False})

# ==========================================
# 1. API PROXY LAYER (Bridge to DB Manager)
# ==========================================
def get_paths(profile): return db.get_paths(profile)
def list_profiles(): return db.list_profiles()
def ensure_profile_structure(name): return db.ensure_profile_structure(name)
def get_all_files_list(profile): return db.get_all_files_list(profile)
def read_file_content(p, f): return db.read_file_content(p, f)
def get_world_state(p): return db.get_world_state(p)
def save_world_state(p, s): return db.save_world_state(p, s)
def get_fragments(p, c): return db.get_fragments(p, c)
def add_fragment(p, n, c, t, tl="", reveal_date=""):
    result = db.add_fragment(p, n, c, t, tl, reveal_date)
    invalidate_retrieval_cache(p)
    return result
def update_fragment(p, i, c, tl=""):
    # Faction versioning — archive old content to Lore only on significant changes
    try:
        old_rows = db.get_fragments(p, "Faction")
        for row in old_rows:
            if row[0] == i:
                old_name = row[1]
                old_content = row[2] or ""
                old_timeline = row[5] or ""

                if old_content.strip() and c.strip():
                    import difflib
                    similarity = difflib.SequenceMatcher(
                        None, old_content.strip(), c.strip()
                    ).ratio()

                    if similarity < 0.90:
                        import datetime
                        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                        archive_name = f"[Archive] {old_name} — {timestamp}"
                        db.add_fragment(p, archive_name, old_content, "Lore", old_timeline)
                        print(f"  [Versioning] Archived '{old_name}' (similarity: {similarity:.2f})")
                    else:
                        print(f"  [Versioning] Minor edit detected for '{old_name}' — skipping archive (similarity: {similarity:.2f})")
                break
    except Exception as e:
        print(f"  [Versioning] Archive failed (non-fatal): {e}")

    result = db.update_fragment(p, i, c, tl)
    invalidate_retrieval_cache(p)
    return result
def update_fragment_metadata(profile, frag_id, new_metadata):
    result = db.update_fragment_metadata(profile, frag_id, new_metadata)
    invalidate_retrieval_cache(profile)
    return result
def update_fragment_type(profile, frag_id, new_type): return db.update_fragment_type(profile, frag_id, new_type)
def update_fragment_reveal_date(profile, frag_id, reveal_date): return db.update_fragment_reveal_date(profile, frag_id, reveal_date)
def update_fragment_known_by(profile, frag_id, known_by): return db.update_fragment_known_by(profile, frag_id, known_by)
def update_fragment_known_versions(profile, frag_id, known_versions): return db.update_fragment_known_versions(profile, frag_id, known_versions)
def get_all_fragments_for_remetadata(profile): return db.get_all_fragments_for_remetadata(profile)
def keyword_search_fragments(profile, query, doc_types=None): return db.keyword_search_fragments(profile, query, doc_types)
def delete_fragment(p, i):
    result = db.delete_fragment(p, i)
    invalidate_retrieval_cache(p)
    return result
def rename_fragment(p, i, n): return db.rename_fragment(p, i, n)
def get_scene_original_draft(profile, filename): return db.get_scene_original_draft(profile, filename)
def get_chat_history(p): return db.get_chat_history(p)
def save_chat_message(p, r, c, session_id=None): return db.save_chat_message(p, r, c, session_id)
def clear_chat_history(p): return db.clear_chat_history(p)
def create_chat_session(profile, name, mode="free"): return db.create_chat_session(profile, name, mode)
def list_chat_sessions(profile): return db.list_chat_sessions(profile)
def get_session_history(profile, session_id): return db.get_session_history(profile, session_id)
def delete_chat_session(profile, session_id): return db.delete_chat_session(profile, session_id)
def rename_chat_session(profile, session_id, name): return db.rename_chat_session(profile, session_id, name)
def lock_chat_message(profile, session_id, message_index, content): return db.lock_chat_message(profile, session_id, message_index, content)
def get_locked_items(profile, session_id): return db.get_locked_items(profile, session_id)
def unlock_chat_message(profile, lock_id): return db.unlock_chat_message(profile, lock_id)
def save_chat_proposal(profile, session_id, content, target_type="", target_id=""): return db.save_chat_proposal(profile, session_id, content, target_type, target_id)
def get_chat_proposals(profile, session_id): return db.get_chat_proposals(profile, session_id)
def update_proposal_status(profile, proposal_id, status): return db.update_proposal_status(profile, proposal_id, status)
def delete_chat_proposal(profile, proposal_id): return db.delete_chat_proposal(profile, proposal_id)
def get_recent_faction_memory(p, f, l=3): return db.get_recent_faction_memory(p, f, l)
def get_all_faction_memories(p): return db.get_all_faction_memories(p)
def update_faction_reaction(p, i, t, f): return db.update_faction_reaction(p, i, t, f)
def delete_faction_reaction(p, i): return db.delete_faction_reaction(p, i)
def save_faction_reaction(p, f, t, s): return db.save_faction_reaction(p, f, t, s)
def add_project(p, n, d, f): return db.add_project(p, n, d, f)
def update_project(p, i, pr, no, nn=None, nd=None): return db.update_project(p, i, pr, no, nn, nd)
def complete_project(p, i, l, t="Fact"): return db.complete_project(p, i, l, t)
def get_fragment_titles_by_ids(profile, ids): return db.get_fragment_titles_by_ids(profile, ids)
def get_generation_log(profile, filename): return db.get_generation_log(profile, filename)
def list_backups(profile): return db.list_backups(profile)
def get_recent_backup_states(profile, count=3): return db.get_recent_backup_states(profile, count)
def restore_backup(profile, filename): return db.restore_backup(profile, filename)
def get_reserved_names(profile): return db.get_reserved_names(profile)
def add_reserved_name(profile, name, note=""): return db.add_reserved_name(profile, name, note)
def delete_reserved_name(profile, name_id): return db.delete_reserved_name(profile, name_id)
def update_reserved_name_note(profile, name_id, note): return db.update_reserved_name_note(profile, name_id, note)

def extract_names_from_scenes(profile_name: str, filenames: List[str]) -> List[str]:
    """
    Runs a lightweight LLM call to extract proper names and usernames
    from the given scene files. Returns a deduplicated list of names.
    """
    paths = db.get_paths(profile_name)
    combined_text = ""
    for filename in filenames:
        filepath = os.path.join(paths['output'], filename)
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                combined_text += f.read() + "\n\n"
        except Exception:
            continue

    if not combined_text.strip():
        return []

    # Cap at 32000 chars to avoid token explosion
    combined_text = combined_text[:32000]

    prompt = f"""Extract every proper name and username from the text below.
Include: character names, usernames, online handles, nicknames, callsigns.
Exclude: place names, organization names, faction names, brand names.
Output ONLY a comma-separated list of names. No explanations, no numbering, no extra text.

TEXT:
{combined_text}"""

    try:
        llm = get_llm(profile_name, "librarian")
        response = llm.invoke([HumanMessage(content=prompt)]).content.strip()
        names = [n.strip() for n in response.split(',') if n.strip()]
        return list(dict.fromkeys(names))  # deduplicate preserving order
    except Exception as e:
        print(f"  [Reserved Names] Extraction failed: {e}")
        return []

def get_reserved_names_block(profile_name: str) -> str:
    """Returns a prompt-ready block of reserved names, or empty string if none."""
    names = db.get_reserved_names(profile_name)
    if not names:
        return ""
    name_list = ", ".join(n['name'] for n in names)
    return f"""*** RESERVED NAMES — DO NOT REUSE ***
The following names and usernames already exist in this story. Do not assign them to new background characters, minor characters, or online personas unless you are deliberately referencing the same individual:
{name_list}
If you need to name a new minor character or generate a username, invent something distinct from this list."""

def get_next_chapter_number(profile_name):
    """Calculates the next available chapter number based on existing files."""
    files = db.get_all_files_list(profile_name)
    max_ch = 0
    for f in files:
        match = re.search(r'Ch(?:apter)?_?(\d+)', f, re.IGNORECASE)
        if match:
            try:
                num = int(match.group(1))
                if num > max_ch: max_ch = num
            except: pass
    return max_ch + 1

# ==========================================
# 2. SETTINGS & AI FACTORY MODULE
# ==========================================

# --- CONFIGURATION PROXIES ---
# These allow the API to read/write settings via the Database Manager

def get_story_settings(profile_name: str) -> dict:
    """Retrieves configuration (Time system, Models, etc.) from DB."""
    return db.get_story_settings(profile_name)

def update_story_setting(profile_name: str, key: str, value: str):
    """Updates a specific configuration key in the DB."""
    db.update_story_setting(profile_name, key, value)

# --- AI MODEL MANAGEMENT ---

class MockResponse:
    def __init__(self, text): self.content = text
class MockLLM:
    def invoke(self, *args, **kwargs): return MockResponse("⚠️ SYSTEM ERROR: API Key missing.")

def list_available_models_all() -> List[str]:
    """
    Dynamically lists ONLY the AI models actually available to your API keys.
    Connects to Google, OpenAI, and Anthropic to fetch real-time lists.
    """
    models = []
    
    # 1. Google Gemini Models
    if GOOGLE_API_KEY:
        try:
            client = new_genai.Client(api_key=GOOGLE_API_KEY)
            for m in client.models.list():
                # Safely get supported methods
                methods = getattr(m, 'supported_generation_methods', [])
                model_name = m.name.replace("models/", "")
                
                # Check for content generation support
                if methods and "generateContent" in methods:
                    models.append(model_name)
                elif "gemini" in model_name.lower() and "embedding" not in model_name.lower():
                    models.append(model_name)
                    
        except Exception as e: 
            print(f"Google Model List Error: {e}")
            # Fallback for Google
            models.extend(["gemini-2.0-flash", "gemini-1.5-pro", "gemini-1.5-flash"])
    
    # 2. OpenAI Models
    if OPENAI_API_KEY:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=OPENAI_API_KEY)
            for m in client.models.list():
                if "gpt" in m.id.lower() or "o1" in m.id.lower():
                    models.append(m.id)
        except Exception as e:
            print(f"OpenAI Model List Error: {e}")

    # 3. Anthropic Models (Dynamic Fetch)
    if ANTHROPIC_API_KEY:
        try:
            import anthropic 
            client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
            
            # Fetch list from API
            for m in client.models.list():
                if "claude" in m.id.lower():
                    models.append(m.id)
                    
        except ImportError:
            print("Anthropic library not installed. Run: pip install anthropic")
        except Exception as e:
            print(f"Anthropic Model List Error: {e}")
            # Fallback if connection fails
            models.extend(["claude-3-5-sonnet-20240620", "claude-3-opus-20240229"])

    return sorted(list(set(models)))

_llm_cache = {} 

def get_llm(profile_name: str, task_type: str = "scene", settings: Optional[dict] = None):
    """
    Factory: Returns the correct LLM client based on the Task and Profile Settings.
    Now 100% provider-agnostic with intelligent fallbacks.
    """
    if settings is None:
        settings = db.get_story_settings(profile_name)
    
    model_map = {
        "scene":     "model_scene",
        "planner":   "model_planner",
        "validator": "model_validator",
        "style":     "model_style",
        "coauthor":  "model_coauthor",
        "reaction":  "model_reaction",
        "warroom":   "model_warroom",
        "librarian": "model_librarian",
    }
    target_key = model_map.get(task_type, "model_coauthor")
    model_name = settings.get(target_key, "")
    
    # --- CHECK CACHE ---
    cache_key = f"{model_name}"
    if model_name and cache_key in _llm_cache:
        return _llm_cache[cache_key]

    client = None

    # --- ATTEMPT TO CREATE REQUESTED CLIENT ---
    if model_name:
        # 1. Google Gemini
        if "gemini" in model_name.lower() and GOOGLE_API_KEY:
            client = ChatGoogleGenerativeAI(model=model_name, google_api_key=GOOGLE_API_KEY)
            
        # 2. OpenAI GPT
        elif ("gpt" in model_name.lower() or "o1" in model_name.lower()) and OPENAI_API_KEY:
            try:
                client = ChatOpenAI(model=model_name, api_key=OPENAI_API_KEY)
            except ImportError:
                print("OpenAI library not installed.")
                
        # 3. Anthropic Claude
        elif "claude" in model_name.lower() and ANTHROPIC_API_KEY and ChatAnthropic:
            client = ChatAnthropic(model=model_name, api_key=ANTHROPIC_API_KEY)

    # --- INTELLIGENT FALLBACK (If requested model failed or no keys matched) ---
    if not client:
        print(f"Warning: Could not load requested model '{model_name}'. Engaging smart fallback...")
        
        # Priority 1: OpenAI
        if OPENAI_API_KEY:
            try:
                client = ChatOpenAI(model="gpt-4o-mini", api_key=OPENAI_API_KEY)
                print("Fallback: Using OpenAI (gpt-4o-mini)")
            except ImportError: pass
            
        # Priority 2: Anthropic
        if not client and ANTHROPIC_API_KEY and ChatAnthropic:
            client = ChatAnthropic(model="claude-3-5-haiku-latest", api_key=ANTHROPIC_API_KEY)
            print("Fallback: Using Anthropic (claude-3-5-haiku-latest)")
            
        # Priority 3: Google
        if not client and GOOGLE_API_KEY:
            client = ChatGoogleGenerativeAI(model="gemini-2.5-flash", google_api_key=GOOGLE_API_KEY)
            print("Fallback: Using Google (gemini-2.5-flash)")

    # --- SAVE TO CACHE & RETURN ---
    if client:
        _llm_cache[cache_key] = client 
        return client
        
    print("CRITICAL ERROR: No valid API keys found in .env file.")
    return MockLLM()

# ==========================================
# 3. SMART RETRIEVAL & HELPERS
# ==========================================

def generate_file_metadata(profile_name: str, content: str) -> str:
    """
    AI Summarizer: Creates a dense, keyword-rich metadata string 
    for the Librarian to use during Smart Retrieval.
    For documents longer than 32000 characters, runs a second pass
    on the second half and merges the results.
    """
    if not content or len(content.strip()) < 50:
        return ""

    def _run_metadata_pass(text_chunk: str) -> str:
        prompt = f"""
    TASK: Generate searchable metadata for the text below.
    
    INSTRUCTION: Read the text and extract the following, in this exact order:
    1. The 15-20 most important proper nouns — prioritize character names, key organizations, and unique locations that someone would use in a search query. Do not list every name mentioned.
    2. The time period covered — a year, year range, or named era.
    3. 3-7 keyword topics describing the key themes, events, and concepts.
    4. A 2-3 sentence summary of what this document establishes.
    
    OUTPUT FORMAT STRICTLY AS (no extra text, no markdown):
    Entities: [Name1, Name2, Name3]
    Period: [e.g. "2016-2021" or "1984" or "Post-War Era" or "The Third Age" or "Before the Collapse" or "Unknown"]
    Topics: [topic1, topic2, topic3]
    Summary: [2-3 sentences describing what this document establishes]
    
    TEXT:
    {text_chunk}
    """
        llm = get_llm(profile_name, "librarian")
        try:
            return llm.invoke([HumanMessage(content=prompt)]).content.strip()
        except Exception as e:
            print(f"Metadata Generation Error: {e}")
            return ""

    # First pass — always runs on first 32000 chars
    first_pass = _run_metadata_pass(content[:32000])

    # Second pass — only runs if document is longer than 32000 chars
    if len(content) <= 32000:
        return first_pass

    print(f"  [Metadata] Document exceeds 32000 chars — running second pass...")
    second_pass = _run_metadata_pass(content[32000:64000])

    if not second_pass:
        return first_pass

    # Merge the two passes
    # Strategy: combine Entities and Topics, keep Period from first pass,
    # concatenate Summaries
    def _extract_field(metadata: str, field: str) -> str:
        for line in metadata.splitlines():
            if line.lower().startswith(field.lower() + ":"):
                return line[len(field) + 1:].strip()
        return ""

    entities_1 = _extract_field(first_pass, "Entities")
    entities_2 = _extract_field(second_pass, "Entities")
    period = _extract_field(first_pass, "Period")
    topics_1 = _extract_field(first_pass, "Topics")
    topics_2 = _extract_field(second_pass, "Topics")
    summary_1 = _extract_field(first_pass, "Summary")
    summary_2 = _extract_field(second_pass, "Summary")

    # Deduplicate entities and topics
    def _merge_lists(a: str, b: str) -> str:
        items_a = [x.strip() for x in a.strip("[]").split(",") if x.strip()]
        items_b = [x.strip() for x in b.strip("[]").split(",") if x.strip()]
        seen = set()
        merged = []
        for item in items_a + items_b:
            if item.lower() not in seen:
                seen.add(item.lower())
                merged.append(item)
        return ", ".join(merged)

    merged_entities = _merge_lists(entities_1, entities_2)
    merged_topics = _merge_lists(topics_1, topics_2)
    merged_summary = " | ".join(filter(None, [summary_1.strip(), summary_2.strip()]))

    return f"Entities: {merged_entities}\nPeriod: {period}\nTopics: {merged_topics}\nSummary: {merged_summary}"

# --- RETRIEVAL CACHE ---
# In-memory cache keyed by query hash. Expires after 30 minutes or on fragment add.
import hashlib
import time

_retrieval_cache: dict = {}
_CACHE_TTL_SECONDS = 1800  # 30 minutes

def _get_cache_key(profile_name: str, user_query: str, doc_types: Optional[List[str]], timeline: str) -> str:
    raw = f"{profile_name}|{user_query}|{sorted(doc_types or [])}|{timeline}"
    return hashlib.md5(raw.encode()).hexdigest()

def invalidate_retrieval_cache(profile_name: str):
    """Clears all cached retrieval results for a profile. Call when fragments are added or modified."""
    keys_to_delete = [k for k in _retrieval_cache if k.startswith(profile_name + "|")]
    for k in keys_to_delete:
        del _retrieval_cache[k]
    if keys_to_delete:
        print(f"  [Cache] Invalidated {len(keys_to_delete)} cached queries for {profile_name}.")

def get_relevant_fragment_ids(profile_name, user_query, doc_types=None, current_timeline="", pov_context=""):
    """
    Scans the 'Table of Contents' (Titles + Metadata) and asks the AI 
    which entries are relevant to the user's query and current timeline.
    Selection cap scales with query length. Low-confidence results are filtered.
    """
    rows = db.get_fragments(profile_name, doc_type=None)
    if not rows: return []

    # Check cache first
    cache_key = f"{profile_name}|{_get_cache_key(profile_name, user_query, doc_types, current_timeline)}"
    cached = _retrieval_cache.get(cache_key)
    if cached:
        entry_time, cached_ids = cached
        if time.time() - entry_time < _CACHE_TTL_SECONDS:
            print(f"  [Librarian] Cache hit — returning {len(cached_ids)} cached fragments.")
            return cached_ids
        else:
            del _retrieval_cache[cache_key]

    if current_timeline:
        rows = _filter_rows_by_timeline(rows, current_timeline)

    # Known-by filter — exclude documents the POV entity doesn't have access to
    if pov_context:
        rows = _filter_rows_by_known_by(rows, pov_context)

    # Format the "Menu" for the AI
    toc_list = []
    for r in rows:
        if doc_types and r[3] not in doc_types:
            continue
        if r[3] == "Reference":
            continue  # Reference fragments injected separately, not via Librarian
        meta_text = ""
        if len(r) > 4 and r[4]:
            clean_meta = r[4].replace('\n', ' ')
            meta_text = f" | {clean_meta}"
        toc_list.append(f"ID: {r[0]} | Title: {r[1]} ({r[3]}){meta_text}")
    
    if not toc_list: return []

    toc_str = "\n".join(toc_list[:150])

    # Scale cap based on named entity count in query
    import re
    # Count capitalised multi-word proper noun clusters as entity references
    entity_matches = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*', user_query)
    unique_entities = set(entity_matches)
    max_items = 20 if len(unique_entities) >= 3 else 15
    print(f"  [Librarian] Query entities detected: {len(unique_entities)} — cap set to {max_items}")

    # --- MULTIVERSE FILTERING INSTRUCTION ---
    timeline_instruction = ""
    if current_timeline:
        timeline_instruction = f"\nNOTE: This scene is set in [{current_timeline}]. The document list above has already been pre-filtered to this timeline and universal documents only.\n"

    prompt = f"""
    ROLE: Database Librarian.
    TASK: Select relevant document IDs based on the user's need and score each one.
    
    *** AVAILABLE DOCUMENTS & METADATA ***
    {toc_str}
    
    *** USER SCENARIO / QUERY ***
    "{user_query}"
    
    *** INSTRUCTION ***
    Analyze the scenario. Identify which documents contain necessary background info based on their Title or Metadata.
    - If the user mentions a specific character, location, or event, check the Entities and Summary fields to find the right file.
    {timeline_instruction}
    - Select ONLY relevant items. Maximum {max_items} items.
    - For each selected item, assign a relevance score from 1-10 where:
      10 = directly named or essential to the scenario
      7-9 = highly relevant background context
      4-6 = marginally relevant, only include if nothing better exists
      1-3 = tangentially related, do not include
    - Do NOT include items scoring below 6.

    OUTPUT FORMAT: JSON list of objects ONLY. Example:
    [{{"id": 1, "score": 9}}, {{"id": 14, "score": 7}}, {{"id": 22, "score": 8}}]
    If nothing is relevant, output: []
    """

    llm = get_llm(profile_name, "librarian")
    try:
        res = llm.invoke([HumanMessage(content=prompt)]).content
        parsed = _extract_json(res)

        if not isinstance(parsed, list):
            return []

        # Handle both old format [1, 2, 3] and new format [{"id": 1, "score": 9}]
        ids = []
        for item in parsed:
            if isinstance(item, int):
                # Fallback: old format with no score, include everything
                ids.append(item)
            elif isinstance(item, dict):
                score = item.get("score", 0)
                if score >= 6:
                    ids.append(item["id"])

        # Store in cache
        _retrieval_cache[cache_key] = (time.time(), ids)
        print(f"  [Librarian] Retrieved {len(ids)} fragments (cap: {max_items}, threshold: 6) — cached.")
        return ids

    except Exception as e:
        print(f"Smart Retrieval Error: {e}")
        return []

def resolve_faction_alias(profile_name, user_input):
    """
    Maps a vague user input (e.g. "The Spies") to a specific Faction Name 
    existing in the database (e.g. "The Guild of Whispers").
    """
    # Use DB Manager to get unique faction names
    existing_factions = db.get_distinct_factions(profile_name)
    
    if not existing_factions: return user_input

    # The "Entity Resolution" Prompt
    prompt = f"""
    TASK: Entity Resolution.
    USER INPUT: "{user_input}"
    KNOWN FACTIONS: {json.dumps(existing_factions)}
    
    INSTRUCTION: Which 'Known Faction' is the user referring to?
    - If it's a nickname (e.g. "The Cops" -> "City Watch"), map it.
    - If it's ambiguous, pick the closest match.
    - If it's a NEW faction not in the list, return "NEW".
    
    OUTPUT: The exact string from the Known Factions list, or "NEW".
    """
    
    llm = get_llm(profile_name, "librarian")
    res = llm.invoke([HumanMessage(content=prompt)]).content.strip()

    res = res.replace('"', '').replace("'", "")
    
    if res == "NEW" or res not in existing_factions:
        return user_input
    return res

def get_last_scenes(profile_name):
    """Retrieves the trailing context (last 3 scenes) for continuity."""
    files = db.get_all_files_list(profile_name)
    
    if not files: return "NO SCENES."
    
    # Files are sorted Newest -> Oldest. We want the chronological order (Old -> New) of the 3 most recent.
    recent_files = files[:3][::-1] 
    
    context = ""
    for f in recent_files:
        content = db.read_file_content(profile_name, f)
        context += f"\n=== PREV: {f} ===\n{content[:3000]}\n"
    return context

def _scene_before_reveal(scene_year: int, scene_date: str, reveal_date: str, profile_name: str) -> bool:
    """
    Asks the LLM whether the reveal_date is still in the future relative to the current scene.
    Handles any calendar system or date format — no Gregorian assumptions.
    Returns True if the spoiler should still be suppressed.
    """
    if not reveal_date:
        return True
    if not scene_date and not scene_year:
        return True

    scene_context = f"{scene_date}, {scene_year}".strip(", ")

    prompt = f"""
    TASK: Determine if a specific date is still in the future relative to the current narrative date.

    Current narrative date: {scene_context}
    Reveal date: {reveal_date}

    Is the reveal date still in the future relative to the current narrative date?
    Answer YES if the reveal date has not happened yet.
    Answer NO if the reveal date has already passed or is the same date.
    Output YES or NO only.
    """

    try:
        llm = get_llm(profile_name, "librarian")
        res = llm.invoke([HumanMessage(content=prompt)]).content.strip().upper()
        return res.startswith("YES")
    except Exception:
        return True

def check_upcoming_spoilers(profile_name: str, scene_year: int, scene_date: str) -> List[dict]:
    """
    Checks if the scene date is dangerously close to any spoiler reveal dates.
    Returns a list of warnings for spoilers that reveal within 30 narrative days.
    Does not suppress anything — purely informational.
    """
    if not scene_year and not scene_date:
        return []

    s_rows = db.get_fragments(profile_name, "Spoiler")
    warnings = []

    for r in s_rows:
        reveal_date = r[6] if len(r) > 6 else ""
        if not reveal_date:
            continue

        content = r[2]
        scene_context = f"{scene_date}, {scene_year}".strip(", ")

        prompt = f"""
        TASK: Determine if a reveal date is imminent relative to the current narrative date.

        Current narrative date: {scene_context}
        Reveal date: {reveal_date}
        Spoiler content: {content[:200]}

        Is the reveal date within approximately 30 narrative days of the current date?
        Answer YES if it is imminent (within ~30 days).
        Answer NO if it is further away or already past.
        Output YES or NO only.
        """

        try:
            llm = get_llm(profile_name, "librarian")
            res = llm.invoke([HumanMessage(content=prompt)]).content.strip().upper()
            if res.startswith("YES"):
                warnings.append({
                    "content": content,
                    "reveal_date": reveal_date,
                    "message": f"Spoiler '{content[:60]}...' reveals on {reveal_date} — close to this scene's date."
                })
        except Exception:
            continue

    return warnings

def check_reveals_passed(profile_name: str, scene_year: int, scene_date: str) -> List[dict]:
    """
    Checks if the scene date is at or past any fragment's reveal_date.
    Returns fragments whose reveal has now passed — user should consider updating Known By.
    Handles any calendar system via LLM comparison.
    """
    if not scene_year and not scene_date:
        return []

    all_frags = db.get_fragments(profile_name, doc_type=None)
    passed = []
    scene_context = f"{scene_date}, {scene_year}".strip(", ")

    for r in all_frags:
        reveal_date = r[6] if len(r) > 6 else ""
        if not reveal_date:
            continue

        name = r[1]
        frag_id = r[0]

        prompt = f"""
        TASK: Determine if a reveal date has been reached or passed relative to the current narrative date.

        Current narrative date: {scene_context}
        Reveal date: {reveal_date}

        Has the reveal date been reached or passed?
        Answer YES if the current date is at or after the reveal date.
        Answer NO if the reveal date is still in the future.
        Output YES or NO only.
        """

        try:
            llm = get_llm(profile_name, "librarian")
            res = llm.invoke([HumanMessage(content=prompt)]).content.strip().upper()
            if res.startswith("YES"):
                passed.append({
                    "id": frag_id,
                    "name": name,
                    "reveal_date": reveal_date
                })
        except Exception:
            continue

    return passed

def audit_scenes_for_spoilers(profile_name: str) -> List[dict]:
    """
    Scans all existing scene files for accidental spoiler leakage.
    Checks each scene against the active spoiler list and flags violations.
    Returns a list of flagged scenes with details.
    """
    # Get all spoilers
    s_rows = db.get_fragments(profile_name, "Spoiler")
    if not s_rows:
        return []

    spoiler_contents = [r[2] for r in s_rows if r[2]]
    if not spoiler_contents:
        return []

    # Get all scene files
    all_files = db.get_all_files_list(profile_name)
    if not all_files:
        return []

    spoiler_list_str = "\n".join(f"- {s}" for s in spoiler_contents)
    flags = []

    for filename in all_files:
        content = db.read_file_content(profile_name, filename)
        if not content or len(content.strip()) < 50:
            continue

        prompt = f"""
        ROLE: Continuity Auditor.
        TASK: Check if this scene accidentally reveals or references any of the listed spoilers.

        *** SPOILERS TO PROTECT (must not appear in scenes yet) ***
        {spoiler_list_str}

        *** SCENE TO CHECK ***
        Filename: {filename}
        Content: {content[:4000]}

        *** INSTRUCTIONS ***
        Read the scene carefully. Check if it directly mentions, hints at, or reveals any of the spoilers above.
        Minor thematic similarity is NOT a violation. Only flag if the scene actually reveals protected information.

        If no spoilers are leaked, output exactly: CLEAN
        If spoilers are leaked, output: LEAKED, then list each violation on a new line.
        Format: LEAKED\n- [spoiler content]: [brief description of how it was leaked in this scene]
        """

        try:
            llm = get_llm(profile_name, "validator")
            res = llm.invoke([HumanMessage(content=prompt)]).content.strip()

            if res.startswith("CLEAN"):
                continue

            if res.startswith("LEAKED"):
                violations = res.replace("LEAKED", "").strip()
                flags.append({
                    "filename": filename,
                    "violations": violations
                })
                print(f"  [Audit] Spoiler leak detected in: {filename}")

        except Exception as e:
            print(f"  [Audit] Error checking {filename}: {e}")
            continue

    return flags

def analyze_world_consequences(profile_name: str, scene_content: str, filename: str) -> dict:
    """
    Analyzes a completed scene and identifies which defined factions and characters
    would plausibly react or be affected. Returns flags for the user to act on.
    Does not generate full reactions — purely advisory.
    """
    # Get defined factions and cast
    world_state = db.get_world_state(profile_name)
    cast = world_state.get("Cast", [])
    char_names = [c.get("Name", "") for c in cast if c.get("Name")]

    faction_rows = db.get_fragments(profile_name, "Faction")
    faction_names = [r[1] for r in faction_rows if r[1]]

    existing_factions = db.get_distinct_factions(profile_name)
    all_factions = list(set(faction_names + (existing_factions or [])))

    if not all_factions and not char_names:
        return {"consequences": [], "message": "No defined factions or characters to analyze against."}

    entities_str = ""
    if all_factions:
        entities_str += f"Known Factions: {', '.join(all_factions)}\n"
    if char_names:
        entities_str += f"Known Characters: {', '.join(char_names[:20])}\n"

    prompt = f"""
    ROLE: World Consequence Analyst.
    TASK: Read this scene and identify which defined factions or characters would plausibly be affected, alerted, or motivated to respond.

    *** DEFINED ENTITIES ***
    {entities_str}

    *** SCENE CONTENT ***
    {scene_content[:4000]}

    *** INSTRUCTIONS ***
    For each entity that would plausibly react to the events in this scene:
    1. Only flag entities that are DEFINED ABOVE — do not invent new factions or characters.
    2. Describe in one sentence WHY they would react and WHAT their likely response direction would be.
    3. Focus on meaningful consequences — not every entity reacts to every scene.
    4. If nothing in this scene would trigger a meaningful response from any defined entity, say so.

    OUTPUT FORMAT: JSON array only.
    Example: [{{"entity": "The Praetorian Guard", "type": "Faction", "reason": "The assassination attempt directly targets their principal — they would mobilize immediately and review security protocols."}}, {{"entity": "Erik Prince", "type": "Character", "reason": "His contractor network was implicated — he would distance himself publicly while privately gathering intelligence."}}]
    If no consequences, output: []
    """

    try:
        llm = get_llm(profile_name, "validator")
        res = llm.invoke([HumanMessage(content=prompt)]).content
        consequences = _extract_json(res)
        if isinstance(consequences, list):
            print(f"  [World Consequences] {len(consequences)} consequence(s) flagged for {filename}")
            return {"consequences": consequences}
        return {"consequences": []}
    except Exception as e:
        print(f"  [World Consequences Error] {e}")
        return {"consequences": []}

def get_global_context(profile_name: str, current_timeline: str = "", scene_year: int = 0, scene_date: str = ""):
    """
    Retrieves the 'Immutable' context layers that must be present in every generation cycle.
    1. Rules: The physics/magic/laws of the world.
    2. Plan: The strategic direction of the story.
    3. Spoilers: Critical secrets to protect — filtered by reveal_date if scene date is known.
    """
    # World Rules
    r_rows = _filter_rows_by_timeline(db.get_fragments(profile_name, "Rulebook"), current_timeline)
    rules = "\n\n".join([r[2] for r in r_rows])

    # Strategic Plan
    p_rows = _filter_rows_by_timeline(db.get_fragments(profile_name, "Plan"), current_timeline)
    plan = p_rows[0][2] if p_rows else "NO PLAN ESTABLISHED."

    # Spoilers — filter by reveal_date if scene date context is available
    s_rows = _filter_rows_by_timeline(db.get_fragments(profile_name, "Spoiler"), current_timeline)
    spoilers = []
    for r in s_rows:
        reveal_date = r[6] if len(r) > 6 else ""
        if not reveal_date:
            # No reveal date — always suppress
            spoilers.append(r[2])
        else:
            # Has reveal date — suppress only if scene hasn't reached it yet
            if _scene_before_reveal(scene_year, scene_date, reveal_date, profile_name):
                spoilers.append(r[2])

    return rules, plan, spoilers

# ==========================================
# 4. SCENE CREATOR ENGINE
# ==========================================

# --- HELPERS ---

def extract_dynamic_spoilers(plan: str, year: int, profile_name: str, settings: Optional[dict] = None, date_str: str = "") -> List[str]:
    """
    Parses future events from the 'Plan' to prevent context leakage into the current narrative.
    """
    if not plan or plan == "NO PLAN ESTABLISHED.":
        return []

    date_context = f"{date_str}, {year}".strip(", ") if date_str else str(year)
    prompt = f"List FUTURE events after {date_context} from: {plan}. OUTPUT: Comma-separated list of event descriptions only."
    llm = get_llm(profile_name, "planner", settings=settings)

    try:
        response = llm.invoke([HumanMessage(content=prompt)]).content
        return [x.strip() for x in response.split(',') if x.strip()]
    except Exception:
        return []
    
def format_state_for_llm(state_dict: dict) -> str:
    """
    Strips UI-specific noise (like coordinates) and formats 
    the World State into highly readable, indented JSON for the AI.
    """
    # Create a deep copy so we don't accidentally delete real data
    import copy
    clean_state = copy.deepcopy(state_dict)
    
    # Strip UI coordinates from Characters
    for char in clean_state.get('Cast', []):
        if 'ui_pos' in char:
            del char['ui_pos']
            
    # Strip UI coordinates from Assets
    for asset in clean_state.get('Assets', []):
        if 'ui_pos' in asset:
            del asset['ui_pos']
            
    # Dump with a 2-space indent for perfect LLM readability
    return json.dumps(clean_state, indent=2)

def _extract_json(text_response: str) -> Any:
    """Safely extracts JSON from an LLM response, even if wrapped in markdown blockquotes."""
    try:
        match = re.search(r'\x60\x60\x60(?:json)?\s*(.*?)\s*\x60\x60\x60', text_response, re.DOTALL | re.IGNORECASE)
        if match:
            return json.loads(match.group(1))
        return json.loads(text_response.strip())
    except Exception as e:
        print(f"JSON Parsing Error: {e} | Raw Text: {text_response[:100]}...")
        return None

def infer_header_data(brief: str, prev_context: str, settings: dict, profile_name: str) -> dict:
    """
    Estimates the narrative date/time for the scene based on recent context using an LLM.
    """
    prompt = f"""
    TASK: Calculate Date/Time/Year.
    BRIEF: {brief}
    CONTEXT END: {prev_context[-500:]}
    DEFAULT TIMEZONE: {settings.get('default_timezone', '')}
    OUTPUT JSON ONLY: {{ "year": 1984, "date": "March 6", "time": "14:00 CST" }}
    """
    llm = get_llm(profile_name, "planner", settings=settings)
    try:
        res = llm.invoke([HumanMessage(content=prompt)]).content
        return _extract_json(res)
    except Exception:
        return {}

def auto_generate_title(profile_name: str, draft_text: str, brief: str) -> str:
    """
    Generates a short, evocative title based on the generated scene content.
    """
    prompt = f"""
    TASK: Create a Title.
    SCENE BRIEF: {brief}
    SCENE CONTENT START: {draft_text[:1000]}...
    
    INSTRUCTION: Generate a short, punchy, dramatic title (max 6 words) for this scene. 
    Examples: "The Red Wedding", "Midnight at the Docks", "Protocol Omega".
    OUTPUT: The title text ONLY. No quotes.
    """
    llm = get_llm(profile_name, "scene")
    try:
        return llm.invoke([HumanMessage(content=prompt)]).content.strip()
    except Exception:
        return "Untitled Scene"

# --- CORE GENERATION LOGIC ---

def _extract_token_usage(response) -> dict:
    """
    Safely extracts token usage from a LangChain AIMessage response object.
    Tries usage_metadata first, falls back to response_metadata.
    """
    try:
        if hasattr(response, 'usage_metadata') and response.usage_metadata:
            u = response.usage_metadata
            return {
                "input": u.get("input_tokens", 0),
                "output": u.get("output_tokens", 0),
                "total": u.get("total_tokens", 0)
            }
        if hasattr(response, 'response_metadata') and response.response_metadata:
            meta = response.response_metadata
            # Gemini format
            if 'usage_metadata' in meta:
                u = meta['usage_metadata']
                return {
                    "input": u.get("prompt_token_count", 0),
                    "output": u.get("candidates_token_count", 0),
                    "total": u.get("total_token_count", 0)
                }
            # OpenAI format via response_metadata
            if 'token_usage' in meta:
                u = meta['token_usage']
                return {
                    "input": u.get("prompt_tokens", 0),
                    "output": u.get("completion_tokens", 0),
                    "total": u.get("total_tokens", 0)
                }
    except Exception:
        pass
    return {"input": 0, "output": 0, "total": 0}

def plan_scene(state: StoryState) -> dict:
    """
    Workflow Node 1: The Director (Planner).
    Reads the heavy context (Lore, Rules, World State) and outputs a detailed Beat Sheet.
    """
    update_generation_status(state['profile_name'], 1, "Planner", "Building scene structure and outline...")
    profile = state['profile_name']
    brief = state['scene_brief']
    current_timeline = state.get('timeline', '').strip()
    
    settings = db.get_story_settings(profile)
    state_tracking = db.get_world_state(profile)
    rules, plan, db_spoilers = get_global_context(
        profile, current_timeline,
        scene_year=state.get('year', 0),
        scene_date=state.get('date_str', '')
    )
    
    print(f"  [Planner] Scanning Knowledge Base for: '{brief[:50]}...'")
    relevant_ids = get_relevant_fragment_ids(
        profile, 
        user_query=brief, 
        doc_types=["Lore", "Fact", "Rulebook", "Scene"],
        current_timeline=current_timeline,
        pov_context=state.get('pov_context', '')
    )
    
    smart_context_str = db.get_content_by_ids(profile, relevant_ids, pov_context=state.get('pov_context', ''))
    if not smart_context_str:
        smart_context_str = "No specific historical records found for this scene."

    dynamic_spoilers = extract_dynamic_spoilers(plan, state['year'], profile, settings=settings, date_str=state.get('date_str', ''))
    all_banned = list(set(db_spoilers + dynamic_spoilers))
    
    use_time_system = settings.get('use_time_system', 'true').lower() == 'true'
    era_display = f"{state['year']}" if (use_time_system and state['year'] > 0) else "Undefined"

    variables_section = ""
    world_vars = state_tracking.get("World Variables", [])
    if world_vars:
        variables_section = "*** WORLD MECHANICS (STRICT) ***\n"
        for v in world_vars:
            variables_section += f"- {v.get('Name', 'Var')}: {v.get('Value', '0')} (RULE: {v.get('Mechanic', '')})\n"

    timeline_section = ""
    if settings.get('use_timelines', 'true').lower() == 'true':
        timelines_list = state_tracking.get("Timelines", [])
        if timelines_list:
            timeline_section = "ACTIVE TIMELINES (MULTIVERSE):\n"
            for t in timelines_list:
                timeline_section += f"- {t.get('Name', 'Unknown')}: {t.get('Description', '')}\n"
    
    if current_timeline:
        timeline_section += f"\n*** CRITICAL: THIS SCENE OCCURS STRICTLY IN TIMELINE: [{current_timeline}] ***\n"
        timeline_section += "Rule: Do NOT reference events, alive/dead statuses, or assets from other timelines unless an explicit crossover is happening.\n"

    prompt = f"""
    ROLE: Lead Story Director & Outliner.
    CURRENT CALENDAR YEAR: {era_display}
    
    *** WORLD LAWS & MECHANICS (STRICT) ***
    {rules}
    {variables_section}
    
    *** STRATEGIC PLAN ***
    {plan}

    *** RELEVANT LORE & CONTEXT (SMART RETRIEVAL) ***
    {smart_context_str}

    *** WORLD STATE & PROJECTS ***
    {format_state_for_llm(state_tracking)}

    {timeline_section}
    
    *** BANNED CONCEPTS (SPOILERS) ***
    [{", ".join(all_banned)}]
    
    *** NARRATIVE CONTEXT (RECENT) ***
    {state['recent_context']}
    
    === MISSION ===
    Create a detailed, 5-to-7 bullet point 'Beat Sheet' (Outline) for the following scene brief.
    Ensure it respects the World Rules, avoids Banned Concepts, and logically follows the Recent Context.
    
    BRIEF: {state['scene_brief']}
    
    OUTPUT: ONLY the bulleted outline. Do not write the prose.
    """
    
    llm = get_llm(profile, "planner", settings=settings)
    _response = llm.invoke([HumanMessage(content=prompt)])
    response = _response.content
    state['_planner_tokens'] = _extract_token_usage(_response)
    
    return {
        "scene_outline": response,
        "banned_words": ", ".join(all_banned),
        "retrieved_ids": relevant_ids
    }

def get_reference_context(profile_name: str) -> tuple[str, str]:
    """
    Fetches Reference fragments split into Style and World sub-types.
    Style references are identified by title starting with [Style] or type tag in metadata.
    World references are everything else in the Reference category.
    Returns (style_block, world_block) as formatted strings.
    """
    ref_rows = db.get_fragments(profile_name, "Reference")
    if not ref_rows:
        return "", ""

    style_refs = []
    world_refs = []

    for r in ref_rows:
        name = r[1] or ""
        content = r[2] or ""
        if not content.strip():
            continue
        # Classify by name prefix or metadata tag
        if name.lower().startswith("[style]") or "[style]" in (r[4] or "").lower():
            style_refs.append(f"--- {name} ---\n{content[:2000]}")
        else:
            world_refs.append(f"--- {name} ---\n{content[:2000]}")

    style_block = "\n\n".join(style_refs) if style_refs else ""
    world_block = "\n\n".join(world_refs) if world_refs else ""

    return style_block, world_block

def _get_faction_pov_block(profile_name: str, character_pov: str) -> str:
    """
    Checks for Main POV factions in the World State.
    If faction POVs are selected, injects collective POV instruction block.
    Works with comma-separated multi-POV context.
    """
    try:
        world_state = db.get_world_state(profile_name)
        all_factions = world_state.get('Factions', [])
        
        # Get faction names from selected POVs (if any)
        selected_povs = [p.strip().lower() for p in character_pov.split(',') if p.strip()] if character_pov else []
        
        if selected_povs:
            # User has selected specific POVs — only include factions that are explicitly selected
            main_pov_factions = [
                f for f in all_factions
                if f.get('Name', '').lower() in selected_povs
                and f.get('Status') != 'Dissolved'
            ]
        else:
            # No POV selected — use Main POV factions automatically
            main_pov_factions = [
                f for f in all_factions
                if f.get('Role') == 'Main POV' and f.get('Status') != 'Dissolved'
            ]

        if not main_pov_factions:
            return ""

        faction_names = ", ".join(f['Name'] for f in main_pov_factions)
        goals = "\n".join(
            f"- {f['Name']}: {f.get('KnownGoals', 'No goals specified')}"
            for f in main_pov_factions
        )
        return f"""*** FACTION POV ***
This scene includes the collective perspective of: {faction_names}
These are institutions, not individuals. Write from a collective institutional voice — multiple members of this faction experience events simultaneously.
Their current goals:
{goals}"""
    except Exception:
        return ""

def draft_scene(state: StoryState) -> dict:
    """
    Workflow Node 2: Narrative Drafting (The Writer).
    Focuses 100% on writing beautiful prose based on the Planner's outline.
    """
    update_generation_status(state['profile_name'], 2, "Drafter", "Writing prose from the approved outline...")
    profile = state['profile_name']
    brief = state['scene_brief']
    chapter = state.get('chapter_num')
    part = state.get('part_num', 1)
    current_timeline = state.get('timeline', '').strip()

    settings = db.get_story_settings(profile)
    state_tracking = db.get_world_state(profile)

    # 1. Retrieve Global Context (Rules, Plan, Spoilers)
    rules, plan, db_spoilers = get_global_context(profile, current_timeline)
    
    # 2. Smart Retrieval (RAG)
    print(f"  [Librarian] Scanning Knowledge Base for: '{brief[:50]}...'")
    relevant_ids = get_relevant_fragment_ids(
        profile, 
        user_query=brief, 
        doc_types=["Lore", "Fact", "Rulebook", "Scene"],
        current_timeline=current_timeline,
        pov_context=state.get('pov_context', '')
    )
    
    smart_context_str = db.get_content_by_ids(profile, relevant_ids, pov_context=state.get('pov_context', ''))
    if not smart_context_str:
        smart_context_str = "No specific historical records found for this scene."

    # Reference Context — style and world texture references
    style_refs, world_refs = get_reference_context(profile)

    # 3. Header & Continuity Logic
    #    If this is Part 2+, fetch the text of previous parts to ensure flow.
    chapter_line = ""
    partition_context = ""

    if chapter is not None:
        chapter_line = f"CHAPTER: {chapter}"
        if part and int(part) > 1:
            chapter_line += f" (PART {part})"
            
            print(f"  [Engine] Fetching continuity for Chapter {chapter}, Part {part}...")
            paths = db.get_paths(profile)
            # Locate previous parts (e.g., Ch01_Part_1...)
            for p in range(1, int(part)):
                pattern = os.path.join(paths['output'], f"Ch{int(chapter):02d}_Part_{p}_*.txt")
                found_files = glob.glob(pattern)
                if found_files:
                    with open(found_files[0], 'r', encoding='utf-8') as f:
                        partition_context += f"\n--- CHAPTER {chapter} PART {p} (PREVIOUS) ---\n{f.read()}\n"

    # 4. Dynamic Spoiler Injection
    #    Prevents the AI from referencing future events defined in the Plan.
    dynamic_spoilers = extract_dynamic_spoilers(plan, state['year'], profile, settings=settings, date_str=state.get('date_str', ''))
    all_banned = list(set(db_spoilers + dynamic_spoilers))
    
    # 5. Chronology & Era Detection
    use_time_system = settings.get('use_time_system', 'true').lower() == 'true'
    header = ""
    era_display = "Undefined (Infer Tech Level from Lore)"
    
    if use_time_system and state['year'] > 0:
        header = f"{state['date_str']}, {state['year']}"
        era_display = f"{state['year']}"
        if state['time_str']: header += f"\n{state['time_str']}"

    # 6. Timeline Logic (Multiverse Support)
    timeline_section = ""
    if settings.get('use_timelines', 'true').lower() == 'true':
        timelines_list = state_tracking.get("Timelines", [])
        if timelines_list:
            timeline_section = "ACTIVE TIMELINES (MULTIVERSE):\n"
            for t in timelines_list:
                timeline_section += f"- {t.get('Name', 'Unknown')}: {t.get('Description', '')}\n"
    
    if current_timeline:
        timeline_section += f"\n*** CRITICAL: THIS SCENE OCCURS STRICTLY IN TIMELINE: [{current_timeline}] ***\n"
        timeline_section += "Rule: Do NOT reference events, alive/dead statuses, or assets from other timelines unless there is an explicit crossover happening right now.\n"
    
    # 7. World Variables (Physics/Mechanics)
    variables_section = ""
    world_vars = state_tracking.get("World Variables", [])
    if world_vars:
        variables_section = "*** WORLD MECHANICS (STRICT) ***\n"
        for v in world_vars:
            variables_section += f"- {v.get('Name', 'Var')}: {v.get('Value', '0')} (RULE: {v.get('Mechanic', '')})\n"

    # 8. Privacy Protocol (Fog of War)
    privacy_protocol = ""
    if state.get('use_fog_of_war', False):
        privacy_protocol = """
        *** PRIVACY & FOG OF WAR PROTOCOL ***
        RULE: Wrap private interactions (whispers, internal thoughts, secure rooms) in [[PRIVATE]] ... [[/PRIVATE]] tags.
        EXAMPLE:
        They stood in the public square. "Everything is fine," he announced loudly.
        [[PRIVATE]]
        Once inside the secure room, he slumped against the door. "We are doomed," he whispered.
        [[/PRIVATE]]
        """

    # 9. Construct Final Prompt (Adaptive Realism)
    #    Note: partition_context is injected to ensure continuity between parts.
    prompt = f"""
    ROLE: Novelist (Third Person Limited).
    CHARACTER: {settings.get('protagonist', 'Protagonist')}.
    {chapter_line}  <-- ONLY INJECT IF EXISTS
    CURRENT CALENDAR YEAR: {era_display}
    
    *** NARRATIVE CONTINUITY (PREVIOUS PARTS) ***
    {partition_context}
    
    *** NARRATIVE LOGIC & TECH-LEVEL (HIERARCHY OF TRUTH) ***
    1. LORE PRIORITY (ABSOLUTE):
       - The 'Story Bible' and 'Rules' are the primary source.
       - If Lore says "Year 407" features Flying Airships, then Airships exist.
       - Do NOT assume "Year 407" means "Real World 407 AD" unless the Lore explicitly confirms it is Earth.
       
    2. DETERMINING THE TECH LEVEL:
       - CHECK LORE FIRST: Scan the Lore below. Does it mention magic, advanced tech, or specific tools? USE THAT.
       - REAL WORLD FALLBACK (Conditional): ONLY if the Lore is SILENT and the setting appears to be Earth, use real-world history for the year {era_display}.
         * Example: Year 1990 + "New York" -> Use Real 1990 Tech (VHS, Landlines).
         * Example: Year 407 + "Kingdom of Asura" -> IGNORE Real 407 AD. Use the Fantasy Logic defined in Rules.
       
    3. REALISM WITHIN CONTEXT: 
       - Once the Tech Level is set (Fantasy or Real), maintain internal consistency.
       - If it's Fantasy, describe the fantasy elements realistically (e.g. the hum of the magic crystal).

    *** WORLD LAWS & MECHANICS (STRICT) ***
    {rules}
    {privacy_protocol}
    {variables_section}
    
    {f"*** STYLE REFERENCE (PROSE RHYTHM & VOICE) ***\n{style_refs}\n" if style_refs else ""}
    {f"*** WORLD TEXTURE REFERENCE (REAL-WORLD MECHANICS) ***\n{world_refs}\n" if world_refs else ""}

    *** STRATEGIC PLAN ***
    {plan}

    *** RELEVANT LORE & CONTEXT (SMART RETRIEVAL) ***
    {smart_context_str}

    *** WORLD STATE & PROJECTS ***
    {format_state_for_llm(state_tracking)}
    
    *** FORMATTING ***
    {header}
    (Start prose below header. NO Title in body).
    
    {timeline_section}
    
    *** BANNED CONCEPTS (SPOILERS) ***
    [{", ".join(all_banned)}]
    
    *** NARRATIVE CONTEXT (RECENT) ***
    {state['recent_context']}
    
    *** APPROVED SCENE OUTLINE ***
    {state.get('scene_outline', 'No outline generated.')}
    
    === MISSION ===
    BRIEF: {state['scene_brief']}
    
    INSTRUCTION: Write the full prose for this scene by expanding the APPROVED SCENE OUTLINE. You are not summarizing events — you are inhabiting them. The lore and world state are not documents to reference. They are the physics of the world you already live in. Write from inside that world, not about it.

    {_get_faction_pov_block(profile, state.get('pov_context', ''))}

    {get_reserved_names_block(profile)}

    *** CRAFT LAWS (NON-NEGOTIABLE) ***

    COMMIT TO THE MOMENT:
    - Every scene must have one clear purpose. If you cannot state in one sentence what this scene must accomplish, find it before writing a word.
    - No atmospheric gesturing. No "the weight of the situation." Either you know what the thing is — say it directly — or it is genuinely unknown — say that and stop. There is no third option.
    - Physical specificity carries tension. Name the object, the gesture, the exact word spoken. Vague impressionism is not mood — it is refusing to commit.

    LORE IS PHYSICS, NOT DIALOGUE:
    - Characters react to the material consequences of events, not to the events as abstract concepts.
    - Characters do not explain the world to each other. They assume shared knowledge. Zero exposition dumping.
    - No character recaps what the reader already knows. If a character references an event, they reference its consequence to them specifically.

    INFORMATION ASYMMETRY:
    - Every character knows only what their specific position in the world allows them to know.
    - A character cannot reference, react to, or be shaped by something that hasn't happened yet or that they have no access to.
    - Characters act from interest, not from narrative convenience.

    AGENCY:
    - Characters pursue goals. They do not wait for the plot to move them.
    - Every character in the scene wants something specific. Their actions follow from that want, not from what the narrative needs to happen next.
    - This applies to antagonists and secondary characters equally. No character exists merely to react to the protagonist.

    INTERIORITY:
    - Interiority must be earned and deployed sparingly. Ask of every moment of internal thought: does the exterior action already carry this? If yes, cut the interior statement.
    - In action sequences: no interiority during the action. Only sensation and decision.
    - In aftermath: interiority compressed into single images or gestures, not paragraphs.

    DIALOGUE:
    - Every line of dialogue must do at least two things simultaneously: advance the scene's purpose AND reveal character, establish relationship dynamic, conceal something, or deflect something.
    - Before finalizing any line: could this line have been written by the author rather than spoken by this character? If yes, rewrite it.
    - Characters do not speak to be helpful to the reader. They speak to get what they want.

    DETAIL:
    - Every specific detail must be load-bearing: it establishes physical reality, reveals character through what they notice, or foreshadows something.
    - Decorative detail that does none of these must be cut.

    PACING:
    - Sentence length controls pace. Shorten sentences as tension rises. Long sentences create density and slow the reader.
    - Read the scene aloud mentally — if the pacing feels wrong in the mouth it will feel wrong on the page.

    FOURTH WALL:
    - Do not reference document titles, file names, or story bible terminology within the scene.
    - Characters exist in their world. The reader observes from outside it.

    CRITIQUE: {state['critique_notes']}
    """
    
    # 10. Execute Generation
    print(f"  [Drafter] Writing prose based on outline...")
    llm = get_llm(profile, "scene", settings=settings)
    _response = llm.invoke([HumanMessage(content=prompt)])
    response = _response.content
    existing = state.get('_drafter_tokens', {"input": 0, "output": 0, "total": 0})
    new_tokens = _extract_token_usage(_response)
    state['_drafter_tokens'] = {
        "input": existing["input"] + new_tokens["input"],
        "output": existing["output"] + new_tokens["output"],
        "total": existing["total"] + new_tokens["total"]
    }
    
    return {
        "current_draft": response, 
        "revision_count": state['revision_count'] + 1, 
        "banned_words": ", ".join(all_banned),
        "retrieved_ids": list(set(state.get('retrieved_ids', []) + relevant_ids))
    }

def critique_scene(state: StoryState) -> dict:
    """
    Workflow Node 3: Validation (The Continuity Editor).
    Rigorously checks the draft against World Rules, State, Banned Words, and the Outline.
    """
    update_generation_status(state['profile_name'], 3, "Validator", "Checking continuity, consistency, and craft...")
    profile = state['profile_name']

    current_timeline = state.get('timeline', '').strip()

    rules, _, _ = get_global_context(profile, current_timeline)
    state_tracking = db.get_world_state(profile)

    timeline_instruction = ""
    if current_timeline:
        timeline_instruction = f"6. TIMELINE CHECK: Did the draft accidentally include people or things that belong to a different timeline? This scene MUST strictly be in [{current_timeline}]."
    
    prompt = f"""
    ROLE: Continuity Editor & Logic Verifier.
    
    You are reviewing a newly drafted scene. Your job is to catch hallucinations, anachronisms, continuity errors, and logic breaks before it is published.

    *** STRICT WORLD RULES ***
    {rules}
    
    *** CURRENT WORLD STATE ***
    {format_state_for_llm(state_tracking)}

    *** BANNED CONCEPTS (DO NOT REVEAL) ***
    [{state['banned_words']}]
    
    *** APPROVED SCENE OUTLINE ***
    {state.get('scene_outline', 'No outline provided.')}

    *** DRAFT TO REVIEW ***
    {state['current_draft']}

    *** INSTRUCTIONS ***
    1. Read the Draft.
    2. SPOILER CHECK: Did the draft accidentally reveal any of the Banned Concepts?
    3. RULE CHECK: Did the draft violate any of the Strict World Rules? (e.g., using magic when magic is forbidden, or modern tech in a medieval setting).
    4. STATE CHECK: Did the draft hallucinate characters who aren't present/alive, or assets the protagonist doesn't own according to the CURRENT WORLD STATE?
    5. OUTLINE CHECK: Did the draft completely ignore the APPROVED SCENE OUTLINE?
    {timeline_instruction}

    6. If the draft is perfectly fine and respects the rules and state, reply EXACTLY with: PASS
    7. If there are errors, reply with: FAIL, followed by a 1-2 sentence explanation of EXACTLY what needs to be fixed.
    """
    
    # Use the 'analysis' model (usually a smarter model like GPT-4o or Claude 3.5 Sonnet)
    llm = get_llm(profile, "validator")
    _response = llm.invoke([HumanMessage(content=prompt)])
    res = _response.content.strip()
    state['_validator_tokens'] = _extract_token_usage(_response)
    
    # 6. Evaluate the AI's critique
    if res.startswith("PASS"):
        return {"is_grounded": True, "critique_notes": ""}
        
    # If it failed, extract the AI's exact notes so the Drafter can fix it.
    notes = res.replace("FAIL", "").strip()
    notes = notes.strip(",").strip(":")
    
    print(f"  [Editor Agent] Rejection! Sending back to Drafter. Reason: {notes}")
    
    return {
        "is_grounded": False, 
        "critique_notes": f"EDITOR FEEDBACK TO FIX: {notes}"
    }

def enforce_style(state: StoryState) -> dict:
    """
    Workflow Node 4: Style Enforcement (The Copy Editor).
    Checks the validated draft against the user's Rulebook for style,
    register, vocabulary, and prose rhythm violations.
    """
    update_generation_status(state['profile_name'], 4, "Style Enforcer", "Refining prose style and voice...")
    profile = state['profile_name']
    settings = db.get_story_settings(profile)

    # Fetch only Rulebook fragments — style rules live there
    current_timeline = state.get('timeline', '').strip()
    rules, _, _ = get_global_context(profile, current_timeline)

    prompt = f"""
    ROLE: Copy Editor & Style Enforcer.

    You are reviewing a scene draft for style compliance only. Continuity and plot have already been verified. Your sole job is to check that the prose matches the established style, register, and vocabulary of this project.

    *** WORLD RULES & STYLE GUIDE ***
    {rules}

    *** DRAFT TO REVIEW ***
    {state['current_draft']}

    *** INSTRUCTIONS ***
    1. Read the Style Guide section of the World Rules above. If no explicit style rules are defined, reply EXACTLY with: PASS
    2. REGISTER CHECK: Does the prose feel consistent with the established tone (e.g. literary, pulpy, clinical, lyrical)? Flag drift.
    3. VOCABULARY CHECK: Are there anachronistic words, modern slang in a period setting, or genre-inappropriate terms?
    4. RHYTHM CHECK: Does the prose style match the established voice, or does it feel generated/generic?

    5. If the draft passes all style checks, reply EXACTLY with: PASS
    6. If there are style violations, reply with: STYLE_FAIL, followed by a 1-2 sentence description of the SPECIFIC violations to fix. Be precise — name the offending words or phrases.
    """

    llm = get_llm(profile, "style", settings=settings)
    _response = llm.invoke([HumanMessage(content=prompt)])
    res = _response.content.strip()
    state['_style_tokens'] = _extract_token_usage(_response)

    if res.startswith("PASS"):
        print(f"  [Style Enforcer] PASS.")
        return {"style_result": "PASS", "style_notes": ""}

    notes = res.replace("STYLE_FAIL", "").strip().strip(",").strip(":")
    print(f"  [Style Enforcer] Violation found. Sending back to Drafter. Reason: {notes}")

    return {
        "style_result": "STYLE_FAIL",
        "style_notes": f"STYLE EDITOR FEEDBACK TO FIX: {notes}",
        "critique_notes": f"STYLE EDITOR FEEDBACK TO FIX: {notes}"
    }

def check_voice_consistency(state: StoryState) -> dict:
    """
    Workflow Node 5: Voice Consistency Check.
    Compares the current draft against the last 3 scenes for character voice consistency.
    """
    update_generation_status(state['profile_name'], 5, "Voice Check", "Verifying character voices and authenticity...")
    profile = state['profile_name']
    settings = db.get_story_settings(profile)

    last_scenes = get_last_scenes(profile)
    if not last_scenes or last_scenes == "":
        print(f"  [Voice Check] No prior scenes found — skipping.")
        return {"voice_result": "PASS", "voice_notes": ""}

    world_state = db.get_world_state(profile)
    cast = world_state.get("Cast", [])
    char_names = [c.get("Name", "") for c in cast if c.get("Name")]
    char_list = ", ".join(char_names[:10]) if char_names else "Unknown"

    prompt = f"""
    ROLE: Character Voice Auditor.
    TASK: Check if the characters in the new draft speak and behave consistently with their voice in recent scenes.

    *** RECENT SCENES (voice reference) ***
    {last_scenes[:6000]}

    *** NEW DRAFT TO CHECK ***
    {state['current_draft'][:4000]}

    *** KNOWN CHARACTERS ***
    {char_list}

    *** INSTRUCTIONS ***
    For each named character who appears in BOTH the recent scenes and the new draft:
    1. Compare their dialogue register — formal/informal, vocabulary level, speech patterns
    2. Compare their emotional baseline — are they acting within their established range?
    3. Compare their behavioral patterns — do their actions match how they've acted before?

    Focus only on clear, specific inconsistencies. Do NOT flag minor variation — characters evolve.
    Flag only if a character sounds like a completely different person.

    If all characters are consistent, output exactly: PASS
    If there are inconsistencies, output: VOICE_FAIL, followed by specific issues per character.
    Example: VOICE_FAIL: John speaks in formal Victorian register here but used street slang in Ch03.
    """

    llm = get_llm(profile, "validator", settings=settings)
    try:
        res = llm.invoke([HumanMessage(content=prompt)]).content.strip()

        if res.startswith("PASS"):
            print(f"  [Voice Check] PASS.")
            return {"voice_result": "PASS", "voice_notes": ""}

        notes = res.replace("VOICE_FAIL", "").strip().strip(",").strip(":")
        print(f"  [Voice Check] Inconsistency found: {notes[:100]}")

        return {
            "voice_result": "VOICE_FAIL",
            "voice_notes": f"VOICE CONSISTENCY FEEDBACK: {notes}",
            "critique_notes": f"VOICE CONSISTENCY FEEDBACK: {notes}"
        }

    except Exception as e:
        print(f"  [Voice Check Error] {e}")
        return {"voice_result": "PASS", "voice_notes": ""}

def get_previous_part_content(profile_name: str, chapter_num: int, part_num: int) -> str:
    """
    Finds and returns the content of the previous part of a chapter.
    Used for part-aware context injection — ensures continuations always
    have the prior part as mandatory context regardless of Librarian results.
    """
    if not chapter_num or not part_num or int(part_num) <= 1:
        return ""

    prev_part = int(part_num) - 1
    paths = db.get_paths(profile_name)
    output_dir = paths['output']

    # Match pattern: Ch{num}_Part_{prev_part}_*
    pattern = os.path.join(output_dir, f"Ch{int(chapter_num):02d}_Part_{prev_part}_*.txt")
    matches = glob.glob(pattern)

    if not matches:
        print(f"  [Part-Aware] No previous part found for Ch{chapter_num:02d} Part {prev_part}.")
        return ""

    # If multiple matches (shouldn't happen but be safe), take the most recent
    matches.sort(key=os.path.getmtime, reverse=True)
    prev_file = matches[0]
    filename = os.path.basename(prev_file)

    try:
        with open(prev_file, 'r', encoding='utf-8') as f:
            content = f.read()
        print(f"  [Part-Aware] Injecting previous part: {filename}")
        return f"\n=== MANDATORY CONTEXT (PREVIOUS PART — DIRECT CONTINUATION) ===\n{content}\n"
    except Exception as e:
        print(f"  [Part-Aware] Failed to read previous part: {e}")
        return ""

def generate_scene(
    profile: str, 
    chapter_num: Optional[int], 
    year: int, 
    date_str: str, 
    time_str: str, 
    title: str, 
    brief: str, 
    context_files: List[str], 
    use_fog_of_war: bool,
    part: int = 1,
    timeline: str = "",
    override_outline: str = "",
    pov_context: str = ""
) -> tuple[str, str]:
    """
    Entry point for the scene generation pipeline.
    Now utilizes a 3-Node Architecture: Planner -> Drafter -> Validator
    """
    # 1. Graph Setup
    workflow = StateGraph(StoryState)
    workflow.add_node("planner", plan_scene)
    workflow.add_node("drafter", draft_scene)
    workflow.add_node("validator", critique_scene)
    workflow.add_node("style_enforcer", enforce_style)
    workflow.add_node("voice_check", check_voice_consistency)

    if override_outline:
        workflow.set_entry_point("drafter")
    else:
        workflow.set_entry_point("planner")
        workflow.add_edge("planner", "drafter")

    workflow.add_edge("drafter", "validator")

    def route_after_validation(state):
        if state['is_grounded']:
            return "style_enforcer"
        if state['revision_count'] > 2:
            print(f"  [WARNING] Revision cap hit on '{state.get('scene_title', 'Unknown')}' — force-passing.")
            return "style_enforcer"
        return "drafter"

    def route_after_style(state):
        if state.get('style_result') == "PASS":
            return "voice_check"
        if state['revision_count'] > 3:
            print(f"  [WARNING] Style cap hit on '{state.get('scene_title', 'Unknown')}' — force-passing.")
            return "voice_check"
        return "drafter"

    def route_after_voice(state):
        if state.get('voice_result') == "PASS":
            return END
        if state['revision_count'] > 4:
            print(f"  [WARNING] Voice cap hit on '{state.get('scene_title', 'Unknown')}' — force-passing.")
            return END
        return "drafter"

    workflow.add_conditional_edges("validator", route_after_validation)
    workflow.add_conditional_edges("style_enforcer", route_after_style)
    workflow.add_conditional_edges("voice_check", route_after_voice)
    app = workflow.compile()
    
    # 2. Context Assembly
    context_str = ""
    if context_files:
        for fname in context_files:
            if fname == "Auto (Last 3 Scenes)": 
                context_str += get_last_scenes(profile)
            else: 
                context_str += f"\n=== CONTEXT: {fname} ===\n{db.read_file_content(profile, fname)}\n"
    else: 
        frags = db.get_fragments(profile, "Lore")
        context_str = f"=== BACKGROUND LORE ===\n{frags[0][2]}" if frags else "NO LORE ESTABLISHED."

    # 2b. Part-Aware Context Injection
    # If this is a continuation (Part > 1), force-inject the previous part
    # as mandatory context before anything else — overrides Librarian results
    prev_part_context = get_previous_part_content(profile, chapter_num, part)
    if prev_part_context:
        context_str = prev_part_context + context_str
        print(f"  [Part-Aware] Previous part prepended to context.")

    settings = db.get_story_settings(profile)
    
    # 3. Heuristic Time Inference
    use_time = settings.get('use_time_system', 'true').lower() == 'true'
    final_year = year
    final_date = date_str
    final_time = time_str
    
    if use_time and (not final_year or not final_date or not final_time):
        inferred = infer_header_data(brief, context_str, settings, profile)
        if not final_year: final_year = inferred.get('year', 1984)
        if not final_date: final_date = inferred.get('date', "Unknown Date")
        if not final_time: final_time = inferred.get('time', "")

    try: final_year = int(final_year)
    except: final_year = 0

    # 4. Chapter Handling
    enable_chapters = str(settings.get('enable_chapters', 'true')).lower() == 'true'
    if not enable_chapters:
        chapter_num = None
    elif chapter_num is None:
        chapter_num = get_next_chapter_number(profile)

    # INITIAL STATE
    temp_title = title if title else "Untitled Processing..."

    initial_input = {
        "profile_name": profile,
        "chapter_num": chapter_num,
        "part_num": part,
        "year": final_year,
        "date_str": final_date,
        "time_str": final_time,
        "scene_title": temp_title,
        "scene_brief": brief,
        "scene_outline": override_outline if override_outline else "",
        "timeline": timeline,
        "recent_context": context_str,
        "revision_count": 0,
        "critique_notes": "",
        "is_grounded": False,
        "current_draft": "",
        "banned_words": "",
        "use_fog_of_war": use_fog_of_war,
        "context_files": context_files,
        "retrieved_ids": [],
        "_planner_tokens": {"input": 0, "output": 0, "total": 0},
        "_drafter_tokens": {"input": 0, "output": 0, "total": 0},
        "_validator_tokens": {"input": 0, "output": 0, "total": 0},
        "_style_tokens": {"input": 0, "output": 0, "total": 0},
        "style_notes": "",
        "style_result": "",
        "voice_notes": "",
        "voice_result": "",
        "pov_context": pov_context if pov_context else settings.get('protagonist', ''),
    }
    
    # 5. Execute AI Loop
    final_state = app.invoke(initial_input)
    clear_generation_status(profile)
    
    # 6. Auto-Title & Persistence
    final_title = title
    if not final_title:
        final_title = auto_generate_title(profile, final_state['current_draft'], brief)

    safe_title = re.sub(r'[\\/*?:"<>|]', "", final_title).replace(" ", "_")

    # Filename Generation
    prefix = ""
    if chapter_num is not None:
        part_suffix = f"_Part_{part}" if part and int(part) > 1 else ""
        prefix = f"Ch{int(chapter_num):02d}{part_suffix}_"

    if use_time:
        safe_date = str(final_date).replace(" ", "-")
        filename = f"{prefix}{final_year}-{safe_date}_{safe_title}.txt"
    else:
        filename = f"{prefix}{safe_title}.txt"
    
    paths = db.get_paths(profile)
    filepath = os.path.join(paths['output'], filename)
    
    # Collision Avoidance
    original_base = filename.replace(".txt", "")
    counter = 1
    while os.path.exists(filepath):
        filename = f"{original_base}_{counter}.txt"
        filepath = os.path.join(paths['output'], filename)
        counter += 1
        
    # [ACTION: WRITE TO FILE]
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(final_state['current_draft'])
        
    # [SYNC TO DATABASE WITH METADATA]
    print(f"  [Engine] Generating metadata for {filename}...")
    metadata = generate_file_metadata(profile, final_state['current_draft'])
    db.upsert_scene(profile, filename, final_state['current_draft'], metadata)

    # [GENERATION LOG]
    retrieved_ids = final_state.get('retrieved_ids', [])
    retrieved_titles = db.get_fragment_titles_by_ids(profile, retrieved_ids)
    validator_result = "PASS" if final_state.get('is_grounded') else "FORCE_PASS"
    style_result = final_state.get('style_result', 'N/A')
    voice_result = final_state.get('voice_result', 'N/A')
    token_usage = {
        "planner": final_state.get('_planner_tokens', {"input": 0, "output": 0, "total": 0}),
        "drafter": final_state.get('_drafter_tokens', {"input": 0, "output": 0, "total": 0}),
        "validator": final_state.get('_validator_tokens', {"input": 0, "output": 0, "total": 0}),
        "style": final_state.get('_style_tokens', {"input": 0, "output": 0, "total": 0}),
        "total": sum(
            v.get("total", 0) for v in [
                final_state.get('_planner_tokens', {}),
                final_state.get('_drafter_tokens', {}),
                final_state.get('_validator_tokens', {}),
                final_state.get('_style_tokens', {})
            ]
        )
    }
    db.save_generation_log(
        profile_name=profile,
        filename=filename,
        brief=brief[:500],
        retrieved_ids=json.dumps(retrieved_ids),
        token_usage=json.dumps(token_usage),
        retrieved_titles=json.dumps(retrieved_titles),
        revision_count=final_state.get('revision_count', 1),
        validator_result=f"{validator_result} | Style: {style_result} | Voice: {voice_result}",
        active_spoilers=final_state.get('banned_words', ''),
        timeline=timeline
    )

    return final_state['current_draft'], filepath

def dry_run_scene(
    profile: str,
    year: int,
    date_str: str,
    time_str: str,
    brief: str,
    context_files: List[str],
    timeline: str = ""
) -> dict:
    """
    Runs only the Planner node — returns the beat sheet outline and retrieved
    document titles without drafting any prose. Zero write operations.
    """
    # --- Context Assembly (mirrors generate_scene) ---
    context_str = ""
    if context_files:
        for fname in context_files:
            if fname == "Auto (Last 3 Scenes)":
                context_str += get_last_scenes(profile)
            else:
                context_str += f"\n=== CONTEXT: {fname} ===\n{db.read_file_content(profile, fname)}\n"
    else:
        frags = db.get_fragments(profile, "Lore")
        context_str = f"=== BACKGROUND LORE ===\n{frags[0][2]}" if frags else "NO LORE ESTABLISHED."

    settings = db.get_story_settings(profile)

    use_time = settings.get('use_time_system', 'true').lower() == 'true'
    final_year = year
    final_date = date_str
    final_time = time_str

    if use_time and (not final_year or not final_date or not final_time):
        inferred = infer_header_data(brief, context_str, settings, profile)
        if not final_year: final_year = inferred.get('year', 0)
        if not final_date: final_date = inferred.get('date', "")
        if not final_time: final_time = inferred.get('time', "")

    try: final_year = int(final_year)
    except: final_year = 0

    # --- Build minimal StoryState for plan_scene ---
    state = {
        "profile_name": profile,
        "chapter_num": None,
        "part_num": 1,
        "year": final_year,
        "date_str": final_date,
        "time_str": final_time,
        "scene_title": "Dry Run",
        "scene_brief": brief,
        "scene_outline": "",
        "timeline": timeline,
        "recent_context": context_str,
        "revision_count": 0,
        "critique_notes": "",
        "style_notes": "",
        "style_result": "",
        "is_grounded": False,
        "current_draft": "",
        "banned_words": "",
        "use_fog_of_war": False,
        "context_files": context_files,
        "retrieved_ids": [],
    }

    # --- Run only the Planner ---
    result = plan_scene(state)

    # --- Resolve retrieved document titles ---
    retrieved_ids = result.get("retrieved_ids", [])
    retrieved_titles = db.get_fragment_titles_by_ids(profile, retrieved_ids)

    return {
        "outline": result.get("scene_outline", ""),
        "retrieved_titles": retrieved_titles,
        "retrieved_ids": retrieved_ids,
        "active_spoilers": result.get("banned_words", ""),
        "inferred_year": final_year,
        "inferred_date": final_date,
        "inferred_time": final_time,
        "upcoming_spoiler_warnings": check_upcoming_spoilers(profile, final_year, final_date),
    }

def split_document_for_ingestion(profile_name: str, content: str, source_name: str,
                                  doc_type: str, timeline: str = "",
                                  threshold: int = 16000) -> List[dict]:
    """
    Splits a large document into focused sub-documents using AI-identified section boundaries.
    Each chunk gets its own metadata. Returns a list of dicts ready for fragment insertion.
    Only splits if content exceeds threshold characters.
    """
    if len(content) <= threshold:
        # Document is small enough — return as single fragment
        metadata = generate_file_metadata(profile_name, content)
        return [{"name": source_name, "content": content, "metadata": metadata}]

    print(f"  [Splitter] Document '{source_name}' is {len(content)} chars — splitting...")

    # Ask AI to identify natural section boundaries
    split_prompt = f"""
    TASK: Identify natural section boundaries in this document for splitting into focused sub-documents.

    DOCUMENT NAME: {source_name}
    DOCUMENT LENGTH: {len(content)} characters

    FIRST 8000 CHARACTERS:
    {content[:8000]}

    MIDDLE SAMPLE (chars 8000-16000):
    {content[8000:16000]}

    INSTRUCTIONS:
    1. Identify 2-5 natural split points where the document shifts to a new major topic.
    2. Each section should be self-contained and cover one coherent subject.
    3. Return character positions (approximate) where splits should occur.
    4. Give each section a short descriptive name.

    OUTPUT FORMAT: JSON array only.
    Example: [{{"name": "Personnel & Command Structure", "start": 0, "end": 8500}}, {{"name": "Operations 2016-2018", "start": 8500, "end": 18000}}]
    """

    llm = get_llm(profile_name, "librarian")
    sections = []

    try:
        res = llm.invoke([HumanMessage(content=split_prompt)]).content
        proposed_sections = _extract_json(res)

        if isinstance(proposed_sections, list) and len(proposed_sections) > 1:
            for s in proposed_sections:
                start = max(0, int(s.get("start", 0)))
                end = min(len(content), int(s.get("end", len(content))))
                chunk = content[start:end].strip()
                if len(chunk) < 100:
                    continue
                section_name = f"{source_name} — {s.get('name', 'Section')}"
                metadata = generate_file_metadata(profile_name, chunk)
                sections.append({
                    "name": section_name,
                    "content": chunk,
                    "metadata": metadata
                })
            if sections:
                print(f"  [Splitter] Split into {len(sections)} sections.")
                return sections
    except Exception as e:
        print(f"  [Splitter] AI split failed: {e}. Falling back to single fragment.")

    # Fallback: return as single fragment if splitting fails
    metadata = generate_file_metadata(profile_name, content)
    return [{"name": source_name, "content": content, "metadata": metadata}]

def bulk_regenerate_metadata(profile_name: str) -> dict:
    """
    Regenerates metadata for every fragment in the profile using the
    current metadata prompt and character cap. Overwrites existing metadata.
    Returns a summary of results.
    """
    fragments = db.get_all_fragments_for_remetadata(profile_name)
    total = len(fragments)
    success = 0
    failed = 0
    skipped = 0

    print(f"  [Bulk Re-metadata] Starting regeneration for {total} fragments...")

    paths = db.get_paths(profile_name)
    conn = sqlite3.connect(paths['db'], timeout=60)
    c = conn.cursor()

    for frag_id, filename, content, frag_type in fragments:
        if not content or len(content.strip()) < 50:
            skipped += 1
            continue
        try:
            new_metadata = generate_file_metadata(profile_name, content)
            if new_metadata:
                c.execute("UPDATE memory_fragments SET metadata = ? WHERE id = ?", (new_metadata, frag_id))
                success += 1
                print(f"  [✓] {filename}")
            else:
                skipped += 1
        except Exception as e:
            print(f"  [✗] {filename}: {e}")
            failed += 1

    conn.commit()
    conn.close()

    print(f"  [Bulk Re-metadata] Done. Success: {success}, Skipped: {skipped}, Failed: {failed}")
    return {"total": total, "success": success, "skipped": skipped, "failed": failed}

def bulk_regenerate_metadata_stream(profile_name: str):
    """
    Generator version of bulk_regenerate_metadata that yields progress events.
    Used by the SSE streaming endpoint.
    Yields dicts with: type, current, total, filename, status
    """
    fragments = db.get_all_fragments_for_remetadata(profile_name)
    total = len(fragments)
    success = 0
    failed = 0
    skipped = 0

    yield {"type": "start", "total": total}

    paths = db.get_paths(profile_name)
    conn = sqlite3.connect(paths['db'], timeout=60)
    c = conn.cursor()

    for i, (frag_id, filename, content, frag_type) in enumerate(fragments):
        current = i + 1
        if not content or len(content.strip()) < 50:
            skipped += 1
            yield {"type": "progress", "current": current, "total": total,
                   "filename": filename, "status": "skipped"}
            continue
        try:
            new_metadata = generate_file_metadata(profile_name, content)
            if new_metadata:
                c.execute("UPDATE memory_fragments SET metadata = ? WHERE id = ?", (new_metadata, frag_id))
                success += 1
                yield {"type": "progress", "current": current, "total": total,
                       "filename": filename, "status": "success"}
            else:
                skipped += 1
                yield {"type": "progress", "current": current, "total": total,
                       "filename": filename, "status": "skipped"}
        except Exception as e:
            failed += 1
            yield {"type": "progress", "current": current, "total": total,
                   "filename": filename, "status": "failed", "error": str(e)}

    conn.commit()
    conn.close()

    yield {"type": "done", "total": total, "success": success,
           "skipped": skipped, "failed": failed}

def save_edited_scene(profile: str, filename: str, content: str) -> tuple[bool, str]:
    """
    Overwrites a scene file with manual edits and updates the database with new metadata.
    """
    try:
        paths = db.get_paths(profile)
        filepath = os.path.join(paths['output'], filename)
        
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
            
        # [Re-generate metadata because the user might have changed important facts]
        print(f"  [Engine] Updating metadata for edited scene {filename}...")
        metadata = generate_file_metadata(profile, content)
        db.upsert_scene(profile, filename, content, metadata)
        
        return True, "Saved successfully."
    except Exception as e:
        return False, str(e)

# --- FILE OPERATIONS (MERGE & DELETE) ---

def merge_specific_files(profile: str, filenames: List[str]) -> str:
    """
    Stitches a user-selected list of files together.
    Archives the source files after a successful merge and updates the DB.
    """
    paths = db.get_paths(profile)
    base_path = paths['output']
    combined_content = ""
    
    # 1. Stitch Content with separator
    for fname in filenames:
        fpath = os.path.join(base_path, fname)
        if os.path.exists(fpath):
            with open(fpath, 'r', encoding='utf-8') as f:
                combined_content += f.read() + "\n\n# # #\n\n"
    
    # 2. Generate New Name
    # Logic: Remove "_Part_X" from the first filename to create the merged title.
    first_name = filenames[0]
    new_name = re.sub(r"_Part_\d+", "", first_name)
    
    if new_name == first_name:
        new_name = "Merged_" + first_name
        
    filepath = os.path.join(base_path, new_name)
    
    # 3. Write new file and Sync to DB
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(combined_content)
        
    # [Sync merged file to DB with fresh metadata]
    print(f"  [Engine] Generating metadata for merged scene {new_name}...")
    metadata = generate_file_metadata(profile, combined_content)
    db.upsert_scene(profile, new_name, combined_content, metadata)
    
    # 4. Archive the original parts
    archive_dir = os.path.join(base_path, "Archive")
    os.makedirs(archive_dir, exist_ok=True)
    
    for fname in filenames:
        source = os.path.join(base_path, fname)
        if os.path.exists(source):
            # Move physical file
            shutil.move(source, os.path.join(archive_dir, fname))
            # [Archive the old scene in the Database]
            db.archive_scene_db(profile, fname)
            
    return new_name


def bulk_delete_files(profile: str, filenames: List[str]) -> int:
    """
    Deletes multiple files in one operation from both disk and Database.
    """
    paths = db.get_paths(profile)
    count = 0
    for fname in filenames:
        fpath = os.path.join(paths['output'], fname)
        if os.path.exists(fpath):
            # 1. Delete physical file
            os.remove(fpath)
            # 2. [Delete from Database]
            db.delete_scene_db(profile, fname)
            count += 1
            
    return count

# ==========================================
# 5. CO-AUTHOR CHAT MODULE
# ==========================================

# ==========================================
# CO-AUTHOR SYSTEM KNOWLEDGE
# ==========================================

CHRONOS_SYSTEM_KNOWLEDGE = """
*** ABOUT CHRONOS STORY DIRECTOR ***
You are the Co-Author embedded in Chronos Story Director — an AI-assisted long-form storytelling system. You can answer questions about how the system works, what features exist, and how to use them effectively.

MODULES:
- Scene Creator: generates prose scenes through a 5-node pipeline (Planner → Drafter → Validator → Style Enforcer → Voice Check). Supports chapters, parts, fog of war, and timeline isolation. Has a Dry Run mode that shows you the outline and retrieved context before spending tokens on full generation.
- Reaction Tool: generates faction reactions to scenes. Uses a preview-before-commit workflow — reactions are shown for review and only appended to the scene file when the user clicks "Append to Scene."
- War Room: strategic consequence simulation. Runs a causality report on a proposed action using world state, faction knowledge, and craft laws (agency, information asymmetry, material consequences). Has optional web search grounding.
- Co-Author Chat (this module): named sessions with four modes — Free Chat, Brainstorm, Scene Repair, Canon/Lore Work. Supports KB attachment via search or file upload, canon locking per message, proposal extraction, contradiction detection, and session summary generation.
- Knowledge Base: Lore, Rules, Plans, Facts, Characters, Factions, Spoilers, Reference tabs. Reference tab holds style references (prefix title with [Style]) and world texture references injected into every generation. The Librarian retrieves relevant documents using AI relevance scoring (threshold 6/10).
- World State Tracker: tracks Cast, Assets, Skills, World Variables, Projects, Relations. Feeds directly into scene generation as structured JSON context. Has AI Batch Analysis to extract state changes from scenes, conflict detection, and backup/restore.
- Network Map: visualizes cast relationships as a force-directed graph with orbital ring layout. Positions are saved per profile.
- Compiler: assembles scenes into manuscript and exports to DOCX.

SETTINGS (accessible via the Settings tab):
- 8 separate model keys: scene, planner, validator, style, coauthor, reaction, warroom, librarian. Each pipeline node can use a different model.
- Time system, chapter system, and multiverse/timeline support are all toggleable.
- Web search available for War Room and Co-Author Chat on Claude and Gemini models only.

PIPELINE NODES (scene generation):
1. Planner — reads lore, rules, world state, and builds a scene outline
2. Drafter — writes prose from the outline, applies craft laws
3. Validator — checks continuity, consistency, and craft against the outline
4. Style Enforcer — checks the draft against your Rulebook for style violations
5. Voice Check — verifies character voice consistency against recent scenes

KNOWLEDGE BASE — HOW RETRIEVAL WORKS:
- The Librarian scores each document 1-10 for relevance to the current scene brief. Only documents scoring 6+ are retrieved.
- Documents need Librarian Metadata (Entities, Period, Topics, Summary fields) to be retrieved accurately. Use bulk re-metadata to generate this automatically.
- Known By field restricts which POV can access a document. Empty = universal.
- Timeline field isolates documents to a specific story timeline.
- Reference documents bypass the Librarian and are always injected.
- reveal_date field triggers post-generation reminders to update Known By when a secret's reveal date is passed.

COMMON ISSUES:
- Document not being retrieved: fill in the Librarian Metadata fields. Use Bulk Re-metadata in the Knowledge Base header to regenerate all metadata at once.
- Reaction Tool not appending: click "Append to Scene" after previewing — it does not auto-save.
- Scene feels disconnected from lore: check that relevant documents have metadata and are not tagged to a different timeline.
- Web search not working in War Room or Co-Author: only Claude and Gemini models support it. Other models fall back gracefully with a note.
- Faction voice feels generic: add a dedicated Faction profile entry in the Factions tab with detailed voice, known information, and blind spots.
- Generation taking a long time: the 5-node pipeline runs 5 sequential LLM calls. The Generate button shows which step is active.
"""


def _get_profile_meta_summary(profile_name: str) -> str:
    """
    Builds a brief summary of the profile's current state for Co-Author meta-awareness.
    Scene count, KB completeness, cast size, world state basics.
    """
    try:
        paths = db.get_paths(profile_name)
        
        # Scene count
        scene_files = db.get_all_files_list(profile_name)
        scene_count = len(scene_files) if scene_files else 0

        # KB entry counts by category
        all_frags = db.get_fragments(profile_name, doc_type=None)
        category_counts = {}
        for r in all_frags:
            cat = r[3] or "Unknown"
            category_counts[cat] = category_counts.get(cat, 0) + 1

        kb_summary_parts = []
        for cat, count in sorted(category_counts.items()):
            kb_summary_parts.append(f"{cat}: {count}")
        kb_summary = ", ".join(kb_summary_parts) if kb_summary_parts else "empty"

        # World state basics
        state = db.get_world_state(profile_name)
        cast_count = len(state.get("Cast", []))
        asset_count = len(state.get("Assets", []))
        has_timelines = len(state.get("Timelines", [])) > 0
        year = state.get("year", 0)

        lines = [
            f"Profile: {profile_name}",
            f"Scenes written: {scene_count}",
            f"Knowledge Base: {kb_summary}",
            f"Cast size: {cast_count} characters",
            f"Assets: {asset_count}",
        ]
        if year and year > 0:
            lines.append(f"Current story year: {year}")
        if has_timelines:
            timeline_names = [t.get("Name", "") for t in state.get("Timelines", [])]
            lines.append(f"Active timelines: {', '.join(timeline_names)}")

        # Completeness hints
        hints = []
        if category_counts.get("Rulebook", 0) == 0:
            hints.append("no Rulebook entries — consider adding world rules and style guidelines")
        if category_counts.get("Faction", 0) == 0:
            hints.append("no Faction profiles — reactions may lack distinct voice")
        if cast_count == 0:
            hints.append("no Cast in World State — add characters to improve scene context")
        if scene_count == 0:
            hints.append("no scenes written yet")

        if hints:
            lines.append(f"Setup notes: {'; '.join(hints)}")

        return "\n".join(lines)

    except Exception as e:
        return f"Profile: {profile_name} (meta-summary unavailable: {e})"

def run_chat_query(profile_name: str, user_input: str, timeline: str = "",
                   mode: str = "free", attached_content: str = "",
                   attached_filename: str = "", session_id: Optional[int] = None) -> str:
    """
    Interacts with the Co-Author persona.
    Mode-aware: brainstorm, scene_repair, canon_work, free
    """
    rules, plan, _ = get_global_context(profile_name, timeline)
    state = db.get_world_state(profile_name)
    settings = db.get_story_settings(profile_name)
    recent_scenes = get_last_scenes(profile_name)
    profile_meta = _get_profile_meta_summary(profile_name)

    relevant_ids = get_relevant_fragment_ids(
        profile_name,
        user_query=user_input,
        doc_types=["Lore", "Fact", "Rulebook", "Scene", "Character", "Faction"],
        current_timeline=timeline,
        pov_context=""
    )
    smart_knowledge = db.get_content_by_ids(profile_name, relevant_ids)
    if not smart_knowledge:
        smart_knowledge = "No specific database records found."

    use_time_system = settings.get('use_time_system', 'true').lower() == 'true'
    era_display = "Undefined"
    if use_time_system and state.get('year', 0) > 0:
        era_display = f"{state['year']}"

    timeline_instruction = ""
    if timeline:
        timeline_instruction = f"\n*** ACTIVE TIMELINE: [{timeline}] ***\nAnswer strictly within this timeline's facts.\n"

    attached_block = ""
    if attached_content:
        attached_block = f"\n*** ATTACHED REFERENCE MATERIAL: {attached_filename} ***\n{attached_content[:8000]}\n"

    # Mode-specific behavior instructions
    mode_instructions = {
        "brainstorm": """
    BEHAVIOR MODE: BRAINSTORM
    - Ask one clarifying question before giving a definitive answer if the query is ambiguous.
    - Propose 2-3 options when multiple valid approaches exist. Let the user choose.
    - Actively flag contradictions with existing lore before answering.
    - After each substantive conclusion, end with: "Ready to lock this as canon?"
    - Think out loud. Show your reasoning, not just your conclusion.
    """,
        "scene_repair": """
    BEHAVIOR MODE: SCENE REPAIR
    - Focus on narrative craft: pacing, dialogue, character voice, tension, continuity.
    - When identifying problems, be specific — name the exact line or passage.
    - Propose concrete rewrites, not abstract advice.
    - Check the attached scene against established character voices in the Knowledge Base.
    - After proposing a fix, ask: "Want me to revise further or is this the direction?"
    """,
        "canon_work": """
    BEHAVIOR MODE: CANON/LORE WORK
    - Prioritize internal consistency above all else.
    - Cross-reference every answer against the Story Bible before stating it.
    - Flag any statement you make that could conflict with existing established facts.
    - When establishing new canon, state it clearly as: "PROPOSED CANON: [statement]"
    - Be conservative — if uncertain, flag it rather than guess.
    """,
        "free": """
    BEHAVIOR MODE: FREE CHAT
    - Answer naturally and helpfully.
    - Use the Story Bible as primary reference.
    - Apply the Hierarchy of Truth: Lore > Real World > Inference.
    """
    }

    behavior = mode_instructions.get(mode, mode_instructions["free"])

    prompt = f"""
    ROLE: Co-Author & Story Collaborator.
    CURRENT YEAR: {era_display}
    {timeline_instruction}
    {behavior}

    *** ABOUT THIS PROJECT ***
    {profile_meta}

    *** ABOUT THIS TOOL ***
    {CHRONOS_SYSTEM_KNOWLEDGE}
    {attached_block}

    *** PRIMARY SOURCE OF TRUTH (STORY BIBLE) ***
    {smart_knowledge}

    *** WORLD RULES ***
    {rules}

    *** FUTURE PLANS ***
    {plan[:3000]}

    *** CURRENT WORLD STATE ***
    {json.dumps(state)}

    *** RECENT NARRATIVE ***
    {recent_scenes}

    *** USER MESSAGE ***
    "{user_input}"

    *** HIERARCHY OF TRUTH ***
    1. Story Bible and Rules = absolute truth
    2. Real world history/science = conditional fallback if lore is silent and setting is Earth-based
    3. Fantasy/alien settings = infer from Rules, never assume Earth
    """

    llm = get_llm(profile_name, "coauthor")
    return llm.invoke([HumanMessage(content=prompt)]).content

def run_chat_query_with_search(profile_name: str, user_input: str, timeline: str = "",
                                mode: str = "free", attached_content: str = "",
                                attached_filename: str = "", session_id: Optional[int] = None) -> str:
    """
    Co-Author Chat with live web search.
    Provider-aware: Claude uses Anthropic search, Gemini uses Google Search grounding.
    """
    rules, plan, _ = get_global_context(profile_name, timeline)
    state = db.get_world_state(profile_name)
    settings = db.get_story_settings(profile_name)
    recent_scenes = get_last_scenes(profile_name)
    profile_meta = _get_profile_meta_summary(profile_name)

    relevant_ids = get_relevant_fragment_ids(
        profile_name,
        user_query=user_input,
        doc_types=["Lore", "Fact", "Rulebook", "Scene", "Character", "Faction"],
        current_timeline=timeline,
        pov_context=""
    )
    smart_knowledge = db.get_content_by_ids(profile_name, relevant_ids)
    if not smart_knowledge:
        smart_knowledge = "No specific database records found."

    use_time_system = settings.get('use_time_system', 'true').lower() == 'true'
    era_display = "Undefined"
    if use_time_system and state.get('year', 0) > 0:
        era_display = f"{state['year']}"

    timeline_instruction = ""
    if timeline:
        timeline_instruction = f"\n*** ACTIVE TIMELINE: [{timeline}] ***\nAnswer strictly within this timeline's facts.\n"

    attached_block = ""
    if attached_content:
        attached_block = f"\n*** ATTACHED REFERENCE MATERIAL: {attached_filename} ***\n{attached_content[:8000]}\n"

    mode_instructions = {
        "brainstorm": "BEHAVIOR MODE: BRAINSTORM — Ask clarifying questions, propose 2-3 options, flag contradictions, end with 'Ready to lock this as canon?'",
        "scene_repair": "BEHAVIOR MODE: SCENE REPAIR — Focus on narrative craft. Name specific problems. Propose concrete rewrites.",
        "canon_work": "BEHAVIOR MODE: CANON/LORE WORK — Prioritize consistency. Flag conflicts. State new canon as 'PROPOSED CANON: [statement]'",
        "free": "BEHAVIOR MODE: FREE CHAT — Answer naturally. Story Bible is primary reference."
    }
    behavior = mode_instructions.get(mode, mode_instructions["free"])

    prompt = f"""
    ROLE: Co-Author & Story Collaborator with real-world research capability.
    CURRENT YEAR: {era_display}
    {timeline_instruction}
    {behavior}

    *** ABOUT THIS PROJECT ***
    {profile_meta}

    *** ABOUT THIS TOOL ***
    {CHRONOS_SYSTEM_KNOWLEDGE}
    {attached_block}

    *** PRIMARY SOURCE OF TRUTH (STORY BIBLE) ***
    {smart_knowledge}

    *** WORLD RULES ***
    {rules}

    *** FUTURE PLANS ***
    {plan[:3000]}

    *** CURRENT WORLD STATE ***
    {json.dumps(state)}

    *** RECENT NARRATIVE ***
    {recent_scenes}

    *** USER MESSAGE ***
    "{user_input}"

    *** RESEARCH CAPABILITY ***
    You have access to web search. Use it when the user asks about real-world facts, historical events, current information, or anything requiring external verification. The Story Bible always takes precedence over search results.

    *** HIERARCHY OF TRUTH ***
    1. Story Bible and Rules = absolute truth
    2. Web search results = real-world grounding when Bible is silent
    3. Fantasy/alien settings = infer from Rules, never assume Earth
    """

    try:
        result, _ = _run_llm_with_web_search(profile_name, prompt, "coauthor")
        return result
    except Exception as e:
        print(f"  [Co-Author Search fallback] {e}")
        return run_chat_query(profile_name, user_input, timeline, mode, attached_content, attached_filename, session_id)

def extract_proposals_from_response(profile_name: str, response_text: str, session_id: int) -> List[dict]:
    """
    Extracts concrete proposals from an AI response.
    Returns structured proposal items for the tracker.
    """
    prompt = f"""
    TASK: Extract concrete proposals from this AI response.
    A proposal is any suggestion to: create a new KB entry, update an existing file, 
    revise a scene, change world state, or establish a new fact.

    RESPONSE TO ANALYZE:
    {response_text[:4000]}

    For each proposal found, output a JSON object with:
    - content: the exact proposed content or change (be specific)
    - target_type: one of "Lore", "Fact", "Plan", "Character", "Faction", "Scene", "WorldState", "New"
    - summary: one sentence describing what this proposal does

    OUTPUT: JSON array only. If no concrete proposals, output: []
    Example: [{{"content": "The Chronos Kernel is a custom OS kernel optimized for I/O throughput", "target_type": "Fact", "summary": "Establishes technical definition of the Chronos Kernel"}}]
    """
    try:
        llm = get_llm(profile_name, "validator")
        res = llm.invoke([HumanMessage(content=prompt)]).content
        proposals = _extract_json(res)
        if isinstance(proposals, list):
            saved = []
            for p in proposals:
                proposal_id = db.save_chat_proposal(
                    profile_name, session_id,
                    p.get("content", ""),
                    p.get("target_type", ""),
                    ""
                )
                saved.append({"id": proposal_id, **p})
            return saved
        return []
    except Exception as e:
        print(f"Proposal extraction error: {e}")
        return []

def check_contradictions_in_content(profile_name: str, content: str, filename: str) -> List[dict]:
    """
    Checks attached content against the Knowledge Base for contradictions.
    Manual trigger only.
    """
    all_frags = db.get_fragments(profile_name, doc_type=None)
    lore_frags = [r for r in all_frags if r[3] in ["Lore", "Fact", "Character", "Faction", "Rulebook"]]
    
    if not lore_frags:
        return []

    kb_summary = "\n".join([f"- [{r[3]}] {r[1]}: {r[4][:200] if r[4] else r[2][:200]}" for r in lore_frags[:50]])

    prompt = f"""
    ROLE: Continuity Auditor.
    TASK: Find contradictions between the attached document and the existing Knowledge Base.

    *** ATTACHED DOCUMENT: {filename} ***
    {content[:5000]}

    *** EXISTING KNOWLEDGE BASE ENTRIES (titles and summaries) ***
    {kb_summary}

    Find direct contradictions — not stylistic differences, but factual conflicts.
    For each contradiction, output:
    - issue: description of the conflict
    - attached_says: what the attached document states
    - kb_says: what the Knowledge Base states
    - kb_entry: which KB entry conflicts

    OUTPUT: JSON array only. If no contradictions, output: []
    """

    try:
        llm = get_llm(profile_name, "validator")
        res = llm.invoke([HumanMessage(content=prompt)]).content
        contradictions = _extract_json(res)
        if isinstance(contradictions, list):
            return contradictions
        return []
    except Exception as e:
        print(f"Contradiction check error: {e}")
        return []

def generate_session_summary(profile_name: str, session_id: int) -> str:
    """
    Generates a compact summary of a chat session.
    """
    history = db.get_session_history(profile_name, session_id)
    locked = db.get_locked_items(profile_name, session_id)

    if not history:
        return "No messages in this session."

    conversation = "\n".join([f"{m['role'].upper()}: {m['content'][:500]}" for m in history[-30:]])
    locked_str = "\n".join([f"- {l['content'][:200]}" for l in locked]) if locked else "None"

    prompt = f"""
    TASK: Summarize this creative writing session concisely.

    *** CONVERSATION ***
    {conversation}

    *** LOCKED CANON ITEMS ***
    {locked_str}

    Produce a summary with these sections:
    1. ESTABLISHED: What was confirmed or canonized
    2. PROPOSED CHANGES: What revisions were suggested (not yet applied)
    3. UNRESOLVED: What questions remain open
    4. FILES TO UPDATE: Which project files likely need updating based on this session

    Be specific and concise. Use bullet points.
    """

    try:
        llm = get_llm(profile_name, "coauthor")
        return llm.invoke([HumanMessage(content=prompt)]).content
    except Exception as e:
        return f"Summary generation failed: {e}"

# ==========================================
# 6. WAR ROOM MODULE
# ==========================================

def _anthropic_with_search(model_name: str, prompt: str) -> str:
    """Runs an Anthropic Claude call with web_search tool enabled."""
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        if "claude" not in model_name.lower():
            model_name = "claude-sonnet-4-20250514"
        response = client.messages.create(
            model=model_name,
            max_tokens=4000,
            tools=[{"type": "web_search_20250305", "name": "web_search"}],
            messages=[{"role": "user", "content": prompt}]
        )
        result = ""
        for block in response.content:
            if hasattr(block, "text"):
                result += block.text
        return result or "No response generated."
    except Exception as e:
        print(f"  [Anthropic Search Error] {e}")
        raise

def _gemini_with_search(model_name: str, prompt: str) -> str:
    """Runs a Gemini call with Google Search grounding enabled."""
    try:
        client = new_genai.Client(api_key=GOOGLE_API_KEY)
        response = client.models.generate_content(
            model=model_name,
            contents=prompt,
            config=new_genai.types.GenerateContentConfig(
                tools=[new_genai.types.Tool(
                    google_search=new_genai.types.GoogleSearch()
                )]
            )
        )
        return response.text
    except Exception as e:
        print(f"  [Gemini Search Error] {e}")
        raise

def _run_llm_with_web_search(profile_name: str, prompt: str, task_type: str) -> tuple[str, bool]:
    """
    Provider-aware web search dispatcher.
    Returns (response_text, search_was_used).
    - Claude models: Anthropic API with web_search_20250305 tool
    - Gemini models: Google Search grounding
    - GPT/others: falls back to standard call with a note added to response
    """
    settings = db.get_story_settings(profile_name)
    model_name = settings.get(f"model_{task_type}", "")

    if "claude" in model_name.lower() and ANTHROPIC_API_KEY:
        print(f"  [Web Search] Using Anthropic search for {model_name}")
        result = _anthropic_with_search(model_name, prompt)
        return result, True

    elif "gemini" in model_name.lower() and GOOGLE_API_KEY:
        print(f"  [Web Search] Using Gemini search grounding for {model_name}")
        result = _gemini_with_search(model_name, prompt)
        return result, True

    else:
        print(f"  [Web Search] Model '{model_name}' does not support web search — running without search.")
        llm = get_llm(profile_name, task_type)
        note = "\n\n[Web search was requested but is not supported by the current model. This response is based on training knowledge only.]\n"
        result = llm.invoke([HumanMessage(content=prompt)]).content
        return result + note, False

def run_war_room_simulation(profile, action_input, timeline=""):
    """
    Executes a Monte Carlo strategic simulation (Smart Retrieval).
    Respects Multiverse/Timeline isolation.
    """
    # 1. Retrieve Global Rules & Plan (Uses local helper wrapping DB)
    rules, plan, _ = get_global_context(profile, timeline)
    
    # 2. Retrieve World State (Direct DB call)
    state = db.get_world_state(profile)
    
    # 3. Get Immediate Tactical Context (Uses local helper)
    recent_history = get_last_scenes(profile)
    
    # 4. Smart Retrieval (Strategic Intelligence)
    print(f"  [War Room] Gathering Intelligence for: '{action_input[:50]}...' in timeline: '{timeline}'")
    relevant_ids = get_relevant_fragment_ids(
        profile, 
        user_query=f"Strategic analysis of: {action_input}", 
        doc_types=["Lore", "Fact", "Rulebook", "Scene"],
        current_timeline=timeline,
        pov_context=""
    )
    
    #    Uses DB Manager for efficient batch content fetching
    smart_intel = db.get_content_by_ids(profile, relevant_ids)
    if not smart_intel:
        smart_intel = "No specific intelligence dossiers found."

    # 5. Construct the Dossier
    #    Note: Uses json.dumps for clean data formatting within the prompt
    # Build faction roster block
    faction_roster = state.get('Factions', [])
    faction_block = ""
    if faction_roster:
        faction_lines = [
            f"- {f['Name']} | Role: {f.get('Role','Neutral')} | Status: {f.get('Status','Unknown')} | Leadership: {f.get('Leadership','Unknown')} | Goals: {f.get('KnownGoals','')}"
            for f in faction_roster
        ]
        faction_block = "*** FACTION ROSTER ***\n" + "\n".join(faction_lines)

    intel_packet = f"""
    *** CURRENT ASSETS & STATUS ***
    Protagonist Status: {json.dumps(state.get('Protagonist Status', {}))}
    Known Cast: {json.dumps(state.get('Cast', []))} 
    
    {faction_block}

    Available Assets: {json.dumps(state.get('Assets', []))}
    Current Skills: {json.dumps(state.get('Skills', []))}
    
    *** IMMEDIATE CONTEXT (RECENT EVENTS) ***
    {recent_history[-4000:]} 
    
    *** RELEVANT KNOWLEDGE (LORE/FACTS) ***
    {smart_intel}
    """
    
    timeline_instruction = ""
    if timeline:
        timeline_instruction = f"\n*** ACTIVE UNIVERSE: [{timeline}] ***\nSimulate consequences strictly within the physics and continuity of this specific timeline. Do not calculate ripple effects into other universes.\n"

    # 6. The "Causality Report" Prompt (Exact Restoration)
    prompt = f"""
    ROLE: Strategic Simulation Engine.
    {timeline_instruction}
    
    *** WORLD RULES & PHYSICS ***
    {rules}
    
    *** CONTEXT PACKET ***
    {intel_packet}
    
    *** CURRENT GOAL ***
    {plan[:2000]}
    
    *** PROPOSED ACTION ***
    "{action_input}"
    
    *** SIMULATION CRAFT LAWS ***

    AGENCY:
    - Every faction and actor in this simulation pursues their own goals. They do not exist to validate or oppose the proposed action — they respond to it from their own interests.
    - Model each faction's reaction from what they want, what they fear, and what they know. Not from what would make a satisfying narrative.

    INFORMATION ASYMMETRY:
    - Each faction knows only what their specific position allows them to know.
    - Do not give factions knowledge of events, capabilities, or intentions they could not have learned through their actual channels.
    - A street-level faction does not know what is happening in a government briefing room unless that information has visibly leaked.

    LORE IS PHYSICS, NOT ABSTRACTION:
    - Simulate material consequences — loss of territory, funding, personnel, credibility, physical infrastructure.
    - Do not describe factions as "displeased" or "alarmed." Describe what they actually do: pull funding, mobilize assets, leak information, accelerate timelines, eliminate threats.

    NO CLEAN RESOLUTION:
    - Real consequences are messy, partial, and contested. Do not resolve tensions neatly.
    - If two factions would realistically have conflicting responses, show both responses without resolving the contradiction.
    - Probability of success is never 0% or 100% unless the World Rules make it physically impossible.

    *** MISSION ***
    Simulate the consequences of this action based on the World Rules.
    Do not write a story. Write a CAUSALITY REPORT.
    
    *** REPORT FORMAT ***
    ## 📊 Simulation Results
    **Probability of Success:** [0-100%]
    
    ### 1. Direct Consequences (Immediate Outcome)
    * [What happens if the action succeeds/fails?]
    * [Cost (Resources, Health, Reputation, or Time)]
    
    ### 2. Second-Order Effects (The Ripple)
    * [Unintended side effects on Relationships/Factions/Environment]
    * [Systemic shifts (Social, Political, Economic, or Magical)] <-- BROADER SCOPE
    
    ### 3. Critical Risks (Blowback)
    * [Who/What reacts negatively?]
    * [Potential catastrophe?]
    
    ### 4. Verdict
    [Go / No-Go recommendation]
    """
    
    # 7. Execution
    llm = get_llm(profile, "warroom")
    return llm.invoke([HumanMessage(content=prompt)]).content

def run_war_room_with_search(profile: str, action_input: str, timeline: str = "") -> str:
    """
    War Room simulation with live web search.
    Provider-aware: Claude uses Anthropic search, Gemini uses Google Search grounding.
    GPT and others fall back gracefully.
    """
    rules, plan, _ = get_global_context(profile, timeline)
    state = db.get_world_state(profile)
    recent_history = get_last_scenes(profile)

    relevant_ids = get_relevant_fragment_ids(
        profile,
        user_query=f"Strategic analysis of: {action_input}",
        doc_types=["Lore", "Fact", "Rulebook", "Scene"],
        current_timeline=timeline,
        pov_context=""
    )
    smart_intel = db.get_content_by_ids(profile, relevant_ids)
    if not smart_intel:
        smart_intel = "No specific intelligence dossiers found."

    faction_roster = state.get('Factions', [])
    faction_block = ""
    if faction_roster:
        faction_lines = [
            f"- {f['Name']} | Role: {f.get('Role','Neutral')} | Status: {f.get('Status','Unknown')} | Leadership: {f.get('Leadership','Unknown')} | Goals: {f.get('KnownGoals','')}"
            for f in faction_roster
        ]
        faction_block = "*** FACTION ROSTER ***\n" + "\n".join(faction_lines)

    intel_packet = f"""
    *** CURRENT ASSETS & STATUS ***
    Protagonist Status: {json.dumps(state.get('Protagonist Status', {}))}
    Known Cast: {json.dumps(state.get('Cast', []))}

    {faction_block}

    Available Assets: {json.dumps(state.get('Assets', []))}
    Current Skills: {json.dumps(state.get('Skills', []))}

    *** IMMEDIATE CONTEXT (RECENT EVENTS) ***
    {recent_history[-4000:]}

    *** RELEVANT KNOWLEDGE (LORE/FACTS) ***
    {smart_intel}
    """

    timeline_instruction = ""
    if timeline:
        timeline_instruction = f"\n*** ACTIVE UNIVERSE: [{timeline}] ***\nSimulate consequences strictly within the physics and continuity of this specific timeline.\n"

    prompt = f"""
    ROLE: Strategic Simulation Engine with real-world research capability.
    {timeline_instruction}

    *** WORLD RULES & PHYSICS ***
    {rules}

    *** CONTEXT PACKET ***
    {intel_packet}

    *** CURRENT GOAL ***
    {plan[:2000]}

    *** PROPOSED ACTION ***
    "{action_input}"

    *** SIMULATION CRAFT LAWS ***

    AGENCY:
    - Every faction and actor in this simulation pursues their own goals. They do not exist to validate or oppose the proposed action — they respond to it from their own interests.
    - Model each faction's reaction from what they want, what they fear, and what they know. Not from what would make a satisfying narrative.

    INFORMATION ASYMMETRY:
    - Each faction knows only what their specific position allows them to know.
    - Do not give factions knowledge of events, capabilities, or intentions they could not have learned through their actual channels.
    - A street-level faction does not know what is happening in a government briefing room unless that information has visibly leaked.

    LORE IS PHYSICS, NOT ABSTRACTION:
    - Simulate material consequences — loss of territory, funding, personnel, credibility, physical infrastructure.
    - Do not describe factions as "displeased" or "alarmed." Describe what they actually do: pull funding, mobilize assets, leak information, accelerate timelines, eliminate threats.

    NO CLEAN RESOLUTION:
    - Real consequences are messy, partial, and contested. Do not resolve tensions neatly.
    - If two factions would realistically have conflicting responses, show both responses without resolving the contradiction.
    - Probability of success is never 0% or 100% unless the World Rules make it physically impossible.

    *** MISSION ***
    Simulate the consequences of this action. You have access to web search — use it to ground your analysis in real-world facts, historical precedents, and current events where relevant. Search for specific data points that would affect the simulation's accuracy.

    *** REPORT FORMAT ***
    ## 📊 Simulation Results
    **Probability of Success:** [0-100%]

    ### 1. Direct Consequences (Immediate Outcome)
    * [What happens if the action succeeds/fails?]
    * [Cost (Resources, Health, Reputation, or Time)]

    ### 2. Second-Order Effects (The Ripple)
    * [Unintended side effects on Relationships/Factions/Environment]
    * [Systemic shifts (Social, Political, Economic, or Magical)]

    ### 3. Critical Risks (Blowback)
    * [Who/What reacts negatively?]
    * [Potential catastrophe?]

    ### 4. Real-World Grounding
    * [Relevant historical precedents or current facts found via search]

    ### 5. Verdict
    [Go / No-Go recommendation]
    """

    try:
        result, _ = _run_llm_with_web_search(profile, prompt, "warroom")
        return result
    except Exception as e:
        print(f"  [War Room Search fallback] {e}")
        return run_war_room_simulation(profile, action_input, timeline)

# ==========================================
# 7. RAG & KNOWLEDGE BASE MODULE
# ==========================================

def get_content_by_ids(profile_name, id_list):
    """
    Retrieves full text content for a specific list of fragment IDs.
    Proxies to the optimized batch fetcher in the database manager.
    """
    return db.get_content_by_ids(profile_name, id_list)

def _filter_rows_by_known_by(rows, pov_context: str):
    """
    Filters fragments based on the known_by field.
    - Empty known_by = Universal, always included
    - "Public" = always included
    - Otherwise: only include if pov_context matches one of the known_by entries
    pov_context is a comma-separated string of entities the current POV has access to.
    """
    if not pov_context:
        return rows

    pov_entities = {e.strip().lower() for e in pov_context.split(',') if e.strip()}
    pov_entities.add('public')  # Public is always accessible

    filtered = []
    for r in rows:
        known_by = r[7] if len(r) > 7 and r[7] else ""
        if not known_by.strip():
            # Empty = Universal — always include
            filtered.append(r)
            continue
        doc_entities = {e.strip().lower() for e in known_by.split(',') if e.strip()}
        # Include if any POV entity matches any doc entity
        if doc_entities & pov_entities:
            filtered.append(r)
    return filtered

def _filter_rows_by_timeline(rows, target_timeline):
    """
    Helper to filter database rows based on the requested timeline.
    Returns fragments that have NO timeline (Universal) or match the target exactly.
    """
    if not target_timeline:
        return rows
        
    filtered = []
    for r in rows:
        # db.get_fragments returns: (id, source_filename, content, type, metadata, timeline)
        row_timeline = r[5] if len(r) > 5 and r[5] else ""
        
        # Keep it if it has no timeline (Universal) or matches the target
        if not row_timeline.strip() or row_timeline.strip().lower() == target_timeline.strip().lower():
            filtered.append(r)
            
    return filtered

def get_global_context(profile_name: str, current_timeline: str = ""):
    """
    Retrieves the 'Immutable' context layers that must be present in every generation cycle.
    1. Rules: The physics/magic/laws of the world.
    2. Plan: The strategic direction of the story.
    3. Spoilers: Critical secrets to protect.
    """
    # Fetch rows from DB Manager: (id, filename, content, type)
    # Index 2 is 'content'
    
    # World Rules
    r_rows = _filter_rows_by_timeline(db.get_fragments(profile_name, "Rulebook"), current_timeline)
    rules = "\n\n".join([r[2] for r in r_rows])
    
    # Strategic Plan
    p_rows = _filter_rows_by_timeline(db.get_fragments(profile_name, "Plan"), current_timeline)
    plan = p_rows[0][2] if p_rows else "NO PLAN ESTABLISHED."
    
    # Spoilers
    s_rows = _filter_rows_by_timeline(db.get_fragments(profile_name, "Spoiler"), current_timeline)
    spoilers = [r[2] for r in s_rows]
    
    return rules, plan, spoilers

def get_full_context_data(profile_name: str, current_timeline: str = ""):
    """
    Retrieves ALL context layers (Lore, Rules, Plans, Facts, Spoilers).
    Used for heavy-duty analysis or deep simulation prompts.
    """
    # Helper to extract content string from rows
    def extract_text(rows): return "\n\n".join([r[2] for r in rows])

    lore = extract_text(_filter_rows_by_timeline(db.get_fragments(profile_name, "Lore"), current_timeline))
    rules = extract_text(_filter_rows_by_timeline(db.get_fragments(profile_name, "Rulebook"), current_timeline))
    
    p_rows = _filter_rows_by_timeline(db.get_fragments(profile_name, "Plan"), current_timeline)
    plan = p_rows[0][2] if p_rows else "NO PLAN."
    
    f_rows = _filter_rows_by_timeline(db.get_fragments(profile_name, "Fact"), current_timeline)
    facts = "\n".join([f"- {r[2]}" for r in f_rows])
    
    s_rows = _filter_rows_by_timeline(db.get_fragments(profile_name, "Spoiler"), current_timeline)
    spoilers = [r[2] for r in s_rows]
    
    return lore, rules, plan, facts, spoilers

def get_initial_lore(profile_name: str, current_timeline: str = "") -> str:
    """Fallback context provider for the initial session if no scenes exist."""
    frags = _filter_rows_by_timeline(db.get_fragments(profile_name, "Lore"), current_timeline)
    if frags:
        return f"=== BACKGROUND LORE ===\n{frags[0][2]}"
    return "NO LORE ESTABLISHED. STARTING FRESH."

def extract_text_from_upload(filename: str, content_bytes: bytes) -> Optional[str]:
    """
    Extracts plain text from an uploaded file.
    Supports .txt, .md, .json, .pdf, and .docx.
    """
    from io import BytesIO

    fname = filename.lower()

    if fname.endswith(".txt") or fname.endswith(".md") or fname.endswith(".json"):
        try:
            return content_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return content_bytes.decode("latin-1", errors="ignore")

    elif fname.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(content_bytes))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n\n".join(p.strip() for p in pages if p.strip())
            return text if text else None
        except Exception as e:
            print(f"PDF extraction error: {e}")
            return None

    elif fname.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(BytesIO(content_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs) if paragraphs else None
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return None

    return None

# --- CRUD PROXIES (Bridge to Database Manager) ---

def get_fragments(profile_name: str, doc_type: Optional[str] = None):
    """Queries memory fragments."""
    return db.get_fragments(profile_name, doc_type)

# ==========================================
# 8. WORLD STATE TRACKER MODULE
# ==========================================

def analyze_state_changes(profile_name, scene_content, timeline=""):
    """
    Executes an LLM analysis of the scene to auto-update the world state (JSON).
    Detects changes in allies, assets, skills, reputation, AND abstract World Variables.
    Multiverse-aware to prevent cross-contamination.
    """
    # Fetch state via DB Manager
    state = db.get_world_state(profile_name)
    
    timeline_instruction = ""
    if timeline:
        timeline_instruction = f"\n*** CRITICAL MULTIVERSE RULE ***\nThese scenes take place STRICTLY in the timeline: [{timeline}]. Only update variables, assets, and character statuses that apply to this specific reality. Do NOT alter facts for other timelines.\n"

    prompt = f"""
    ROLE: World State Database Manager.
    
    TASK: Analyze the TEXT CONTENT (Scene, Lore, or Plan) and update the JSON STATE.
    {timeline_instruction}
    
    *** CURRENT STATE ***
    {format_state_for_llm(state)}
    
    *** NARRATIVE SCENE ***
    {scene_content}
    
    *** UPDATE INSTRUCTIONS ***
    1. TIME & DATES (CRITICAL EXECUTION):
       - Compare any detected years against the 'Current_Year' ({state.get('Current_Year', 'Unknown')}).
       - RULE A: "NARRATIVE ONLY": ONLY update 'Current_Year' if the NARRATIVE VOICE confirms the story has actually reached that time.
         * YES: "The year 2030 finally arrived..." (Update to 2030)
         * YES: "Ten years passed..." (Add 10 years)
         * NO: Dialogue references (e.g. "I will be done in 2038") -> IGNORE.
         * NO: Future Plans/Visions (e.g. "He foresaw the crash of 2029") -> IGNORE.
       - RULE B: "FORWARD ONLY": Never update to a year older than Current_Year (Flashback protection).
       - RULE C: BIRTH YEAR: If explicitly mentioned as a fact (e.g. "Born in 1984"), update 'Protagonist Status' -> 'Birth_Year'.

    2. WORLD VARIABLES (CRITICAL):
       - Review the 'World Variables' list in the State.
       - Based on the scene's events, strictly Apply the "Mechanic/Rule" defined for each variable.
       - Example: If 'Federal Heat' rule says "violence increases this", and scene has violence, increase the Value.
       - Output the UPDATED list of variables.

    3. CAST & ROSTER: 
       - Update the 'Cast' list. 
       - For existing characters (match by Name), update 'Role' or 'Tags' if their status changes.
       - If a new MAJOR character appears, add them to 'Cast' (Role='Support').
       - Update 'Loyalty' numbers based on interactions.

    4. ASSETS: Add new resources/locations gained. Mark lost assets as "Destroyed".

    5. SKILLS: Add new skills learned.

    6. ALIASES & REPUTATION:
       - Look for new titles, nicknames, or reputations bestowed upon the protagonist by the public or other characters.
       - Example: If they conquer a city, add "Conqueror of [City]".
       - Example: If they fix the economy, add "The Architect".
       - MERGE these into the existing 'Aliases' string in 'Protagonist Status' (comma-separated).

    7. PROJECTS & RESEARCH:
       - Review the 'Projects' list.
       - If the narrative describes significant work, breakthroughs, or testing related to a project, INCREASE its "Progress" (0-100).
       - Small effort: +5-10%. Major breakthrough: +20-50%. Completion: Set to 100%.
       - If the project is ruined/destroyed, reduce progress.
    
    CRITICAL OUTPUT RULE: 
    You must return the COMPLETE JSON STATE object, including all unchanged fields (Protagonist, Lore, etc.). 
    DO NOT return a partial update. The output must be the full, valid JSON structure.
    """
    
    llm = get_llm(profile_name, "warroom")
    try:
        res = llm.invoke([HumanMessage(content=prompt)]).content
        new_state = _extract_json(res)
        
        if new_state:
            if "Protagonist Status" not in new_state:
                state.update(new_state)
                final_state = state
            else:
                final_state = new_state

            conflicts = detect_state_conflicts(profile_name, final_state)
            return {"proposed_state": final_state, "conflicts": conflicts}

        return {"proposed_state": state, "conflicts": []}

    except Exception as e:
        print(f"Analysis Error: {e}")
        return {"proposed_state": state, "conflicts": []}

def detect_state_conflicts(profile_name: str, proposed_state: dict) -> List[dict]:
    """
    Compares a proposed world state against the last 3 backups.
    Returns a list of flagged contradictions for user review.
    """
    recent_states = db.get_recent_backup_states(profile_name, count=3)
    if not recent_states:
        return []

    # Build a compact summary of recent states for comparison
    # We only send Cast, Assets, Variables, and Current_Year — not the full state
    def compact(state):
        return {
            "Current_Year": state.get("Current_Year", "Unknown"),
            "Cast": [
                {
                    "Name": c.get("Name"),
                    "Role": c.get("Role"),
                    "Loyalty": c.get("Loyalty"),
                    "Tags": c.get("Tags", [])
                }
                for c in state.get("Cast", [])
            ],
            "Assets": [
                {"Name": a.get("Name"), "Status": a.get("Status", "Active")}
                for a in state.get("Assets", [])
            ],
            "World Variables": [
                {"Name": v.get("Name"), "Value": v.get("Value")}
                for v in state.get("World Variables", [])
            ]
        }

    recent_compact = [compact(s) for s in recent_states]
    proposed_compact = compact(proposed_state)

    prompt = f"""
    ROLE: Continuity Auditor.
    TASK: Compare a proposed world state update against recent saved states and flag contradictions.

    *** RECENT SAVED STATES (newest first) ***
    {json.dumps(recent_compact, indent=2)}

    *** PROPOSED NEW STATE ***
    {json.dumps(proposed_compact, indent=2)}

    *** INSTRUCTIONS ***
    Check for contradictions between the proposed state and the recent saves:
    1. CHARACTER: Loyalty changed by more than 25 points with no gradual trend in recent saves
    2. CHARACTER: Role changed (e.g. Support -> Antagonist) without a matching trend
    3. CHARACTER: A character present in recent saves is now missing entirely
    4. ASSET: An asset marked Active in recent saves is now missing or Destroyed without a clear trend
    5. YEAR: Current_Year moved backwards (flashback protection)
    6. VARIABLE: A world variable changed dramatically in a single jump

    For each contradiction found, output a JSON object with:
    - field: the field name (e.g. "Cast.John Smith.Loyalty")
    - old_value: what it was in the most recent save
    - new_value: what the proposed state says
    - reason: one sentence explaining why this is suspicious

    OUTPUT: JSON array only. If no contradictions found, output: []
    Example: [{{"field": "Cast.John.Loyalty", "old_value": 80, "new_value": 20, "reason": "Loyalty dropped 60 points in a single update with no gradual trend."}}]
    """

    try:
        llm = get_llm(profile_name, "validator")
        res = llm.invoke([HumanMessage(content=prompt)]).content
        conflicts = _extract_json(res)
        if isinstance(conflicts, list):
            return conflicts
        return []
    except Exception as e:
        print(f"Conflict detection error: {e}")
        return []

# ==========================================
# 9. NETWORK MAP
# ==========================================

def generate_network_graph(profile: str):
    """
    Constructs the node/edge graph from the new 'Cast' Roster.
    Supports Multi-POV (Constellation) layouts.
    """
    state = db.get_world_state(profile)
    cast = state.get("Cast", [])
    assets = state.get("Assets", [])
    
    nodes = []
    edges = []
    
    # 1. Build Character Nodes
    for char in cast:
        is_pov = char.get("Role") == "POV"
        
        # Determine Visual Category for styling
        category = "Ally"
        if is_pov: category = "Protagonist" 
        elif char.get("Role") == "Antagonist": category = "Enemy"
        
        nodes.append({
            "id": char["id"],
            "type": "customNode",
            "data": { 
                "label": char["Name"],
                "icon": char.get("Icon", "Neutral"),
                "category": category,
                "role": char.get("Role", "Support"),
                "orbit": char.get("Orbit", None),
                "timeline": char.get("Timeline", ""),
                "ring": char.get("Ring", "")
            },
            # Default pos (Frontend will auto-arrange)
            "position": char.get("ui_pos", {"x": 0, "y": 0}) 
        })

        # 2. Build Character Edges (Links)
        for link in char.get("Links", []):
            target_id = link["targetId"]
            
            # Create a unique sorted ID for the edge so A->B and B->A don't create two lines
            edge_id = f"e-{sorted([char['id'], target_id])[0]}-{sorted([char['id'], target_id])[1]}"
            
            # Only add if not already in edges list
            if not any(e['id'] == edge_id for e in edges):
                edges.append({
                    "id": edge_id,
                    "source": char["id"],
                    "target": target_id,
                    "label": link["type"]
                })

    # 3. Build Asset Nodes
    # Assets orbit the Main POV (or the first found POV) by default
    first_pov = next((c for c in cast if c.get("Role") == "POV"), None)
    
    for i, asset in enumerate(assets):
        asset_id = f"asset_{i}"
        nodes.append({
            "id": asset_id,
            "type": "customNode",
            "data": {
                "label": asset.get("Asset", "Item"),
                "icon": asset.get("Icon", "Resource"),
                "category": "Asset"
            },
            "position": asset.get("ui_pos", {"x": 0, "y": 0})
        })
        
        # Create 'Owns' link
        if first_pov:
            edges.append({
                "id": f"e-{first_pov['id']}-{asset_id}",
                "source": first_pov['id'],
                "target": asset_id,
                "label": "Owns"
            })

    return {"nodes": nodes, "edges": edges}

def update_character_link_label(profile: str, source_id: str, target_id: str, new_label: str):
    """
    Updates the relationship label between two cast members in the World State.
    """
    state = db.get_world_state(profile)
    cast = state.get("Cast", [])
    
    for char in cast:
        if char["id"] == source_id:
            links = char.get("Links", [])
            for link in links:
                if link["targetId"] == target_id:
                    link["type"] = new_label
                    break
            break
    
    state["Cast"] = cast
    db.save_world_state(profile, state)

# ==========================================
# 10. REACTION TOOL & FACTION LOGIC
# ==========================================

def save_faction_reaction(profile_name, faction, text, scene_name):
    """Logs a raw reaction to the database via DB Manager."""
    db.save_faction_reaction(profile_name, faction, text, scene_name)

def get_recent_faction_memory(profile_name, faction, limit=3):
    """Retrieves the last few raw reactions via DB Manager."""
    return db.get_recent_faction_memory(profile_name, faction, limit)

def undo_last_reaction_text(profile_name, filename, faction):
    """
    Removes the last appended reaction for a specific faction from the text file.
    Safeguard: Only deletes if the file explicitly ends with a reaction block for this faction.
    """
    paths = db.get_paths(profile_name)
    filepath = os.path.join(paths['output'], filename)
    
    if not os.path.exists(filepath): return False, "File not found."
    
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    header_marker = f">>> REACTION: {faction}"
    
    if header_marker not in content:
        return False, "No reaction text found in file."

    parts = content.rsplit(header_marker, 1)
    
    if len(parts) < 2:
        return False, "Could not isolate reaction block."

    clean_content = parts[0].rstrip()

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(clean_content)
        
    db.upsert_scene(profile_name, filename, clean_content)
    
    return True, "Reaction text stripped from file."

def delete_last_faction_reaction(profile, faction):
    db.delete_last_faction_reaction(profile, faction)

def preview_reaction_for_scene(profile_name, filename, faction, public_only=False, format_style="Standard", custom_instructions="", timeline=""):
    """
    Generates a reaction preview without appending to the file or saving to the database.
    Returns the reaction text for user review before committing.
    """
    true_faction = resolve_faction_alias(profile_name, faction)
    rules, plan, _ = get_global_context(profile_name, timeline)
    state = db.get_world_state(profile_name)
    content = db.read_file_content(profile_name, filename)
    settings = db.get_story_settings(profile_name)

    use_time = settings.get('use_time_system', 'true').lower() == 'true'
    era_display = "Undefined (Infer Tech Level from Lore)"
    if use_time and state.get('year', 0) > 0:
        era_display = f"{state['year']}"

    past_reactions = db.get_recent_faction_memory(profile_name, true_faction)

    query = f"Faction '{true_faction}' reacting to scene content: {content[:3000]}..."
    relevant_ids = get_relevant_fragment_ids(
        profile_name,
        user_query=query,
        doc_types=["Lore", "Fact", "Rulebook", "Scene", "Faction"],
        current_timeline=timeline,
        pov_context=""
    )
    smart_facts = db.get_content_by_ids(profile_name, relevant_ids, pov_context=true_faction)
    style_refs, world_refs = get_reference_context(profile_name)

    if public_only:
        pattern = r"\[\[PRIVATE\]\].*?\[\[/PRIVATE\]\]"
        content = re.sub(pattern, "[...INTERNAL/PRIVATE SCENE REDACTED...]", content, flags=re.DOTALL | re.IGNORECASE)
        content = re.sub(r"\[PRIVATE:.*?\]", "[REDACTED]", content)

    knowledge_instr = (
        "You are reading the unredacted scene. HOWEVER, act strictly as the Target Faction. "
        "DO NOT reference internal thoughts of others unless the provided Rules/Lore explicitly grant telepathic abilities. "
        "Otherwise, react ONLY to observable actions."
    )
    if public_only:
        knowledge_instr = (
            "CRITICAL: You are an EXTERNAL OBSERVER. Private interactions have been REDACTED. "
            "You DO NOT know what happened in the redacted sections. Do NOT guess accurately."
        )

    timeline_instruction = ""
    if timeline:
        timeline_instruction = f"\n*** ACTIVE UNIVERSE: [{timeline}] ***\nReact strictly based on the history, tech, and facts of this specific timeline.\n"

    faction_profile = ""
    faction_rows = db.get_fragments(profile_name, "Faction")
    for row in faction_rows:
        if true_faction.lower() in row[1].lower() or true_faction.lower() in row[2].lower():
            faction_profile = row[2]
            break

    prompt = f"""
    ROLE: Narrative Simulator (Grounded in History & State).
    TARGET FACTION: {true_faction}
    CURRENT YEAR/ERA: {era_display}
    {timeline_instruction}

    *** WORLD STATE & DATA ***
    Character Roster: {json.dumps(state.get('Cast', []))}

    *** RELEVANT INTELLIGENCE (SMART RETRIEVAL) ***
    {smart_facts}

    *** FACTION PROFILE (PRIMARY VOICE REFERENCE) ***
    {faction_profile if faction_profile else "No dedicated profile found. Infer voice from past reactions and world state."}

    {f"*** STYLE REFERENCE ***\n{style_refs}\n" if style_refs else ""}
    {f"*** WORLD TEXTURE REFERENCE ***\n{world_refs}\n" if world_refs else ""}

    *** PAST REACTIONS (VOICE EVOLUTION) ***
    {past_reactions if past_reactions else "No prior reactions on record. This is the first reaction from this faction."}

    *** MISSION ***
    Write a reaction to the SCENE provided below. You are not summarizing what happened. You are generating the specific, partial, self-interested response of this specific entity at this specific moment based on what they know and what they want.

    *** FORMAT ADAPTATION PROTOCOL (CRITICAL) ***
    Requested Format: "{format_style}"

    INSTRUCTION: You must check if the Requested Format exists in the Current Era ({era_display}).
    1. IF COMPATIBLE: Use the format as requested (e.g. "Newspaper" in 1920).
    2. IF ANACHRONISTIC: Transmute the format to the closest era-appropriate equivalent.

    *** HIERARCHY OF TRUTH ***
    1. LORE PRIORITY: If Lore says "Telepathy exists," then "Mental Chat" is a valid format.
    2. REALISM: Use real-world logic for the Era to determine how news travels.

    {get_reserved_names_block(profile_name)}

    *** REACTION CRAFT LAWS (NON-NEGOTIABLE) ***

    VOICE AUTHENTICITY:
    - Before writing any line, ask: could this sentence have been written in a neutral analytical register rather than spoken by this specific entity? If yes, rewrite it until the answer is no.
    - The test is not whether the content is correct for the character — it is whether the specific words, rhythm, and construction belong to that character and not to the author.
    - Voice containment is absolute. Never cross-contaminate registers. A politician cares about polling, donors, and institutional power. A street-level operative thinks in logistics and threat assessment. A bureaucrat speaks in passive constructions and hedged language. Each voice is a closed system.

    INFORMATION ASYMMETRY:
    - This faction knows only what their specific position allows them to know.
    - They do not have omniscient knowledge of events they could not have witnessed or learned through their specific channels.
    - They react to what breached their awareness, not to the full event as the reader knows it.

    LORE IS PHYSICS, NOT DIALOGUE:
    - React to the material consequences of the scene — loss of funding, shift in power, new threat, physical destruction.
    - Do not recite or reference the scene's events as abstract concepts. React to what those events mean for this faction's specific interests.

    NO RESOLUTION (DEFAULT — overridable via Additional Instructions):
    - Unless instructed otherwise, real reactions do not conclude in consensus or neat summation.
    - If there is internal disagreement, it remains unresolved at the end. Characters repeat themselves with increasing intensity. They talk past each other. They run out of steam without concluding.
    - A character saying "fair point" and the argument ending is almost always false. Let disagreements remain disagreements.

    DRIFT:
    - Real conversations do not stay on topic. Allow natural drift — one thing becomes adjacent, then further, then circles back or doesn't.
    - Emotional register shifts without announcement: serious analysis, then a one-word dismissal, then something unexpectedly personal, then back.
    - Do not maintain a consistent tone across the entire reaction.

    NO RECAP:
    - Do not start by summarizing what happened. Jump directly into the reaction already in progress.
    - No character recaps events the others already know. They argue about those events, dispute them, feel things about them.

    BOTH BAN:
    - The word "both" used as a conclusion is banned. Noticing that two things are simultaneously true is not a thought — it is the beginning of a thought that was not finished.
    - If there is a tension, the character argues with it, feels something about it, or goes somewhere new with it.

    FOURTH WALL:
    - Do not reference document titles, file names, or story bible terminology.
    - Characters respond to what exists in their world.

    *** ADDITIONAL INSTRUCTIONS ***
    {custom_instructions if custom_instructions else "Follow standard personality and lore. To direct a specific outcome — a concession, a pivot, a change in position — state it here explicitly."}

    *** KNOWLEDGE CONSTRAINTS ***
    {knowledge_instr}

    *** SCENE CONTEXT ***
    {content}
    """

    llm = get_llm(profile_name, "reaction")
    res = llm.invoke([HumanMessage(content=prompt)]).content

    if past_reactions and faction_profile:
        consistency_prompt = f"""
        ROLE: Continuity Editor.
        TASK: Check if the new faction reaction is consistent with this faction's established voice and known information.

        *** FACTION PROFILE (GROUND TRUTH) ***
        {faction_profile}

        *** PAST REACTIONS (VOICE HISTORY) ***
        {past_reactions}

        *** NEW REACTION TO CHECK ***
        {res}

        INSTRUCTION: Check for:
        1. Voice inconsistency — does the tone, register, or personality match the profile?
        2. Knowledge contradiction — does the faction reference information it shouldn't know yet?
        3. Behavioral contradiction — does the faction act against its established character?

        If consistent, output: CONSISTENT
        If inconsistent, output: INCONSISTENT: [one sentence describing the specific contradiction]
        """
        try:
            consistency_llm = get_llm(profile_name, "validator")
            consistency_check = consistency_llm.invoke([HumanMessage(content=consistency_prompt)]).content.strip()
            if consistency_check.startswith("INCONSISTENT"):
                res = res + f"\n\n⚠️ CONSISTENCY WARNING: {consistency_check.replace('INCONSISTENT:', '').strip()}"
        except Exception as e:
            print(f"  [Consistency Check Error] {e}")

    if "REFUSAL" in res:
        return False, res

    return True, res

def commit_reaction_to_scene(profile_name: str, filename: str, faction: str,
                              reaction_text: str, format_style: str = "Standard"):
    """
    Appends a previewed reaction to the scene file and saves it to the database.
    Called only after user confirms the preview.
    """
    true_faction = resolve_faction_alias(profile_name, faction)
    db.save_faction_reaction(profile_name, true_faction, reaction_text, filename)

    paths = db.get_paths(profile_name)
    clean_style = format_style.split("->")[-1].strip()
    header = f"\n\n>>> REACTION: {true_faction} | {clean_style} <<<\n"
    full_filepath = os.path.join(paths['output'], filename)

    with open(full_filepath, "a", encoding="utf-8") as f:
        f.write(header + reaction_text + "\n")

    with open(full_filepath, "r", encoding="utf-8") as f:
        full_updated_content = f.read()

    db.upsert_scene(profile_name, filename, full_updated_content)
    return True

# ==========================================
# 11. COMPILER MODULE
# ==========================================

def compile_manuscript(profile_name, files):
    """Compiles selected files into a single manuscript."""
    return "\n***\n".join([db.read_file_content(profile_name, f) for f in files])

def compile_docx_manuscript(profile_name: str, selected_files: List[str],
                             font_name: str = "Times New Roman",
                             font_size: int = 12,
                             margin_cm: float = 2.54) -> bytes:
    """
    Compiles selected scene files into a formatted .docx Word document.
    """
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --- PAGE MARGINS ---
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)

    # --- PARAGRAPH STYLE HELPER ---
    def set_body_format(para):
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt = para.paragraph_format
        fmt.first_line_indent = Inches(0.5)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = Pt(font_size * 1.5)

    # --- TITLE PAGE ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(200)
    run = title_para.add_run(profile_name.replace("_", " "))
    run.bold = True
    run.font.name = font_name
    run.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(20)
    sub_run = sub.add_run("Generated by Chronos Story Director")
    sub_run.font.name = font_name
    sub_run.font.size = Pt(11)
    sub_run.italic = True

    doc.add_page_break()

    # --- CHAPTER LOOP ---
    for filename in selected_files:
        raw_content = db.read_file_content(profile_name, filename)

        # Clean system tags
        raw_content = raw_content.replace("[[PRIVATE]]", "").replace("[[/PRIVATE]]", "")

        # Smart chapter title
        base_name = filename.replace(".txt", "")
        chapter_prefix = ""
        match = re.search(r'(Ch\d+|Chapter_\d+)', base_name, re.IGNORECASE)
        if match:
            try:
                num = int(re.search(r'\d+', match.group(0)).group(0))
                chapter_prefix = f"Chapter {num}: "
            except: pass
        clean_parts = [p for p in base_name.split("_") if not re.match(r'(Ch\d+|Chapter|\d{4})', p)]
        chapter_title = f"{chapter_prefix}{' '.join(clean_parts)}"

        # Chapter heading
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(36)
        heading_run = heading.add_run(chapter_title)
        heading_run.bold = True
        heading_run.font.name = font_name
        heading_run.font.size = Pt(font_size + 6)

        # --- BODY: join wrapped lines, split on blank lines into paragraphs ---
        # A blank line = intentional paragraph break
        # A non-blank line followed by another non-blank line = same paragraph (line wrap)
        raw_paragraphs = re.split(r'\n\s*\n', raw_content.strip())

        # Timestamp pattern: lines that look like dates/times at the start of a scene
        timestamp_pattern = re.compile(
            r'^[\w]+ \d+\w*,?\s*\d{4}|^\d{2}:\d{2}|^Year:|^Date:|^Time:',
            re.IGNORECASE
        )

        first_para = True
        for block in raw_paragraphs:
            lines = [l.strip() for l in block.splitlines() if l.strip()]
            if not lines:
                continue

            # Check if the first line of this block looks like a timestamp header
            # If so, split it out as its own styled paragraph before the rest
            if first_para and lines and timestamp_pattern.match(lines[0]):
                # Render timestamp line(s) as italic centered header
                header_text = ' '.join(lines[:2]) if len(lines) > 1 and timestamp_pattern.match(lines[1]) else lines[0]
                remaining_lines = lines[2:] if len(lines) > 1 and timestamp_pattern.match(lines[1]) else lines[1:]

                ts_para = doc.add_paragraph()
                ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ts_para.paragraph_format.space_before = Pt(0)
                ts_para.paragraph_format.space_after = Pt(18)
                ts_run = ts_para.add_run(header_text)
                ts_run.italic = True
                ts_run.font.name = font_name
                ts_run.font.size = Pt(font_size - 1)

                if not remaining_lines:
                    first_para = False
                    continue
                lines = remaining_lines

            text = ' '.join(lines)
            if not text:
                continue

            para = doc.add_paragraph()
            set_body_format(para)

            if first_para:
                para.paragraph_format.first_line_indent = Pt(0)
                first_para = False

            para_run = para.add_run(text)
            para_run.font.name = font_name
            para_run.font.size = Pt(font_size)

        doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def compile_formatted_manuscript(profile_name: str, selected_files: List[str]) -> Dict[str, bytes]:
    """
    Compiles selected scene files into professional PDF and EPUB formats.
    Performs 'Typesetting' cleanups:
    - Removes [[PRIVATE]] tags (keeps content).
    - Formats 'Reactions' as proper Interludes/Dossiers.
    - Sanitizes smart quotes/dashes for PDF compatibility.
    """
    
    # --- HELPER: TEXT CLEANER ---
    def clean_manuscript_text(text):
        # Remove System Tags (Privacy)
        text = text.replace("[[PRIVATE]]", "").replace("[[/PRIVATE]]", "")
        
        # Sanitize Smart Characters for PDF (Latin-1 safe)
        replacements = {
            '\u201c': '"', '\u201d': '"',  # Smart double quotes
            '\u2018': "'", '\u2019': "'",  # Smart single quotes
            '\u2013': '-', '\u2014': '--', # Dashes
            '\u2026': '...',               # Ellipsis
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
            
        return text

    # --- HELPER: REACTION FORMATTER ---
    def format_reaction_blocks(text):
        """
        Converts raw '>>> REACTION' blocks into stylish 'Interludes'.
        Removes the '✨ Custom' debug lines and metadata clutter.
        """
        # Split the text into the Main Scene and appended Reactions
        parts = re.split(r'>>> REACTION:', text)
        
        # Part 0 is the main story
        final_text = clean_manuscript_text(parts[0].strip())
        
        # Process any reactions (Parts 1+)
        if len(parts) > 1:
            for raw_reaction in parts[1:]:
                # Regex to parse the header: "Faction | Type <<<"
                header_match = re.search(r'\s*(.*?) \|\s*(.*?) <<<', raw_reaction)
                
                if header_match:
                    faction = header_match.group(1).strip()
                    r_type = header_match.group(2).strip()
                    
                    # Remove the header line AND the following "✨" line if it exists
                    body = re.sub(r'\s*(.*?) \|\s*(.*?) <<<(\n✨.*)?', '', raw_reaction, count=1).strip()
                    
                    # Kill the "Category -> Format" line
                    body = re.sub(r'^.* -> .*$', '', body, flags=re.MULTILINE)
                    
                    # Kill the "PARTIES:" line entirely
                    body = re.sub(r'^\s*(\*\*)?PARTIES:.*$', '', body, flags=re.MULTILINE)
                    
                    # Remove double stars (**Name**) if they exist in the body
                    body = re.sub(r'\*\*(.*?)\*\*', r'\1', body)

                    # Clean up extra empty lines created by the deletions
                    body = re.sub(r'\n{3,}', '\n\n', body).strip()

                    # Sanitize body text
                    body = clean_manuscript_text(body)
                    
                    # PDF/EPUB visual separator
                    final_text += "\n\n" + ("*" * 20) + "\n\n" 
                    final_text += f"INTERLUDE: {faction.upper()}\n"
                    final_text += f"Type: {r_type}\n"
                    final_text += ("-" * 20) + "\n\n"
                    final_text += body
                else:
                    # Fallback if regex fails (just clean and append)
                    final_text += "\n\n***\n\n" + clean_manuscript_text(raw_reaction)
                    
        return final_text

    # --- DATA AGGREGATION ---
    chapters = []
    for filename in selected_files:
        # Use DB Manager to read file content
        raw_content = db.read_file_content(profile_name, filename)
        
        # Apply the Typesetting Logic
        formatted_body = format_reaction_blocks(raw_content)

        # Smart Title Logic (Ch04 -> Chapter 4)
        base_name = filename.replace(".txt", "")
        chapter_prefix = ""
        match = re.search(r'(Ch\d+|Chapter_\d+)', base_name, re.IGNORECASE)
        if match:
            try:
                num = int(re.search(r'\d+', match.group(0)).group(0))
                chapter_prefix = f"Chapter {num}: "
            except: pass

        clean_parts = [p for p in base_name.split("_") if not re.match(r'(Ch\d+|Chapter|\d{4})', p)]
        raw_title = " ".join(clean_parts)
        final_title = f"{chapter_prefix}{raw_title}"

        chapters.append({"title": final_title, "body": formatted_body})

    results = {"pdf": None, "epub": None}

    # --- PDF PIPELINE (fpdf2) ---
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Front Matter
        pdf.add_page()
        pdf.set_font("Times", "B", 24)
        pdf.cell(0, 60, f"Story Profile: {profile_name}", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Times", "", 12)
        pdf.cell(0, 10, "Generated by Chronos Story Director", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.add_page()

        # Chapter Loop
        for chap in chapters:
            pdf.set_font("Times", "B", 16)
            pdf.cell(0, 10, chap['title'], new_x="LMARGIN", new_y="NEXT")
            pdf.ln(5)
            
            pdf.set_font("Times", "", 12)
            # Encode/Decode to handle Latin-1 limitations of standard FPDF
            safe_body = chap['body'].encode('latin-1', 'ignore').decode('latin-1')
            pdf.multi_cell(0, 6, safe_body)
            pdf.ln(10) 
            pdf.add_page()

        results["pdf"] = bytes(pdf.output())
    except Exception as e:
        print(f"PDF Generation Error: {e}")

    # --- EPUB PIPELINE (EbookLib) ---
    try:
        book = epub.EpubBook()
        book.set_identifier(profile_name)
        book.set_title(profile_name)
        book.set_language('en')

        epub_chapters = []
        for i, chap in enumerate(chapters):
            c = epub.EpubHtml(title=chap['title'], file_name=f'chap_{i}.xhtml', lang='en')
            html_body = chap['body'].replace("\n", "<br/>")
            html_body = html_body.replace("********************", "<hr/>")
            
            c.content = f"<h1>{chap['title']}</h1><p>{html_body}</p>"
            book.add_item(c)
            epub_chapters.append(c)

        book.toc = (epub_chapters)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        style = 'body { font-family: serif; } h1 { text-align: center; }'
        nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
        book.add_item(nav_css)
        book.spine = ['nav'] + epub_chapters
        
        buffer = BytesIO()
        epub.write_epub(buffer, book, {})
        buffer.seek(0)
        results["epub"] = buffer.getvalue()
    except Exception as e:
        print(f"EPUB Generation Error: {e}")

    return results